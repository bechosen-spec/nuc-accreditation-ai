# import re
# import json
# import sqlite3
# import hashlib
# from datetime import datetime

# import streamlit as st
# import pandas as pd
# from openai import OpenAI

# from sklearn.model_selection import train_test_split
# from sklearn.preprocessing import LabelEncoder
# from sklearn.linear_model import LogisticRegression
# from sklearn.tree import DecisionTreeClassifier
# from sklearn.ensemble import RandomForestClassifier, GradientBoostingClassifier
# from sklearn.ensemble import VotingClassifier

# # =========================================================
# # PAGE CONFIG
# # =========================================================
# st.set_page_config(
#     page_title="NUC Accreditation Prediction System",
#     page_icon="🎓",
#     layout="wide",
#     initial_sidebar_state="expanded"
# )

# # =========================================================
# # CUSTOM STYLING
# # =========================================================
# st.markdown("""
# <style>
#     .block-container {
#         padding-top: 1.2rem;
#         padding-bottom: 1.5rem;
#     }

#     .hero-card {
#         padding: 1.2rem 1.4rem;
#         border-radius: 18px;
#         border: 1px solid rgba(120,120,120,0.15);
#         background: linear-gradient(135deg, rgba(59,130,246,0.10), rgba(99,102,241,0.08));
#         margin-bottom: 1rem;
#     }

#     .score-card {
#         padding: 1rem;
#         border-radius: 16px;
#         border: 1px solid rgba(120,120,120,0.12);
#         background: rgba(34,197,94,0.06);
#         text-align: center;
#         min-height: 120px;
#     }

#     .report-box {
#         padding: 1.2rem;
#         border-radius: 16px;
#         border: 1px solid rgba(120,120,120,0.15);
#         background: rgba(99,102,241,0.05);
#         line-height: 1.7;
#     }

#     div[data-testid="stMetric"] {
#         background: rgba(255,255,255,0.03);
#         border: 1px solid rgba(120,120,120,0.10);
#         padding: 0.8rem;
#         border-radius: 14px;
#     }

#     .stButton > button {
#         border-radius: 12px;
#         padding: 0.65rem 1rem;
#         font-weight: 600;
#     }
# </style>
# """, unsafe_allow_html=True)

# st.markdown("""
# <div class="hero-card">
#     <h1 style="margin-bottom:0;">🎓 NUC Accreditation Prediction System</h1>
# </div>
# """, unsafe_allow_html=True)

# # =========================================================
# # DATABASE
# # =========================================================
# DB_PATH = "nuc_app.db"


# def get_conn():
#     return sqlite3.connect(DB_PATH, check_same_thread=False)


# def init_db():
#     conn = get_conn()
#     cur = conn.cursor()

#     cur.execute("""
#         CREATE TABLE IF NOT EXISTS users (
#             id INTEGER PRIMARY KEY AUTOINCREMENT,
#             full_name TEXT NOT NULL,
#             email TEXT NOT NULL UNIQUE,
#             password_hash TEXT NOT NULL,
#             role TEXT NOT NULL DEFAULT 'user',
#             created_at TEXT NOT NULL
#         )
#     """)

#     cur.execute("""
#         CREATE TABLE IF NOT EXISTS assessments (
#             id INTEGER PRIMARY KEY AUTOINCREMENT,
#             user_id INTEGER NOT NULL,
#             institution_name TEXT NOT NULL,
#             discipline TEXT NOT NULL,
#             programme_name TEXT NOT NULL,
#             self_study_score REAL NOT NULL,
#             predicted_status TEXT NOT NULL,
#             actual_ratio REAL,
#             core_pct REAL,
#             phd_pct REAL,
#             created_at TEXT NOT NULL,
#             assessment_payload TEXT NOT NULL,
#             report_text TEXT,
#             FOREIGN KEY(user_id) REFERENCES users(id)
#         )
#     """)

#     conn.commit()
#     conn.close()


# init_db()


# def hash_password(password: str) -> str:
#     return hashlib.sha256(password.encode("utf-8")).hexdigest()


# def create_user(full_name: str, email: str, password: str):
#     conn = get_conn()
#     cur = conn.cursor()
#     try:
#         cur.execute(
#             """
#             INSERT INTO users (full_name, email, password_hash, role, created_at)
#             VALUES (?, ?, ?, ?, ?)
#             """,
#             (
#                 full_name,
#                 email.strip().lower(),
#                 hash_password(password),
#                 "user",
#                 datetime.utcnow().isoformat()
#             )
#         )
#         conn.commit()
#         return True, "Account creation is successful."
#     except sqlite3.IntegrityError:
#         return False, "An account with that email already exists."
#     finally:
#         conn.close()


# def authenticate_user(email: str, password: str):
#     conn = get_conn()
#     cur = conn.cursor()
#     cur.execute(
#         """
#         SELECT id, full_name, email, role
#         FROM users
#         WHERE email = ? AND password_hash = ?
#         """,
#         (email.strip().lower(), hash_password(password))
#     )
#     row = cur.fetchone()
#     conn.close()
#     return row


# def save_assessment(user_id: int, institution: str, discipline: str, programme: str,
#                     self_study_score: float, predicted_status: str, actual_ratio: float,
#                     core_pct: float, phd_pct: float, assessment_payload: str, report_text: str):
#     conn = get_conn()
#     cur = conn.cursor()
#     cur.execute(
#         """
#         INSERT INTO assessments (
#             user_id, institution_name, discipline, programme_name,
#             self_study_score, predicted_status, actual_ratio, core_pct, phd_pct,
#             created_at, assessment_payload, report_text
#         )
#         VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
#         """,
#         (
#             user_id, institution, discipline, programme,
#             self_study_score, predicted_status, actual_ratio, core_pct, phd_pct,
#             datetime.utcnow().isoformat(), assessment_payload, report_text
#         )
#     )
#     conn.commit()
#     conn.close()


# def get_user_assessments(user_id: int):
#     conn = get_conn()
#     df = pd.read_sql_query(
#         """
#         SELECT id, institution_name, discipline, programme_name,
#                self_study_score, predicted_status, actual_ratio,
#                core_pct, phd_pct, created_at
#         FROM assessments
#         WHERE user_id = ?
#         ORDER BY datetime(created_at) DESC
#         """,
#         conn,
#         params=(user_id,)
#     )
#     conn.close()
#     return df


# def get_user_report_history(user_id: int):
#     conn = get_conn()
#     df = pd.read_sql_query(
#         """
#         SELECT id, institution_name, discipline, programme_name,
#                predicted_status, self_study_score, created_at, report_text
#         FROM assessments
#         WHERE user_id = ?
#         ORDER BY datetime(created_at) DESC
#         """,
#         conn,
#         params=(user_id,)
#     )
#     conn.close()
#     return df


# def get_all_users():
#     conn = get_conn()
#     df = pd.read_sql_query(
#         """
#         SELECT id, full_name, email, role, created_at
#         FROM users
#         ORDER BY datetime(created_at) DESC
#         """,
#         conn
#     )
#     conn.close()
#     return df


# def get_all_assessments():
#     conn = get_conn()
#     df = pd.read_sql_query(
#         """
#         SELECT id, user_id, institution_name, discipline, programme_name,
#                self_study_score, predicted_status, created_at
#         FROM assessments
#         ORDER BY datetime(created_at) DESC
#         """,
#         conn
#     )
#     conn.close()
#     return df

# # =========================================================
# # MODEL TRAINING
# # =========================================================
# @st.cache_resource
# def train_model():
#     df = pd.read_csv("nuc_dataset_22150_questionnaire_full.csv")

#     target = "actual_accreditation_status"
#     drop_cols = [
#         "programme_id",
#         "institution_name",
#         "programme_name",
#         "self_study_score",
#         "academic_score",
#         "staffing_score",
#         "physical_facilities_score",
#         "library_score",
#         "funding_score",
#         "research_score",
#         target,
#     ]

#     X = df.drop(columns=drop_cols)
#     y = df[target]

#     le = LabelEncoder()
#     y = le.fit_transform(y)

#     X = pd.get_dummies(X, drop_first=True)
#     training_columns = X.columns

#     X_train, _, y_train, _ = train_test_split(
#         X, y, test_size=0.2, random_state=42, stratify=y
#     )

#     lr = LogisticRegression(max_iter=1000)
#     dt = DecisionTreeClassifier(random_state=42)
#     rf = RandomForestClassifier(n_estimators=150, random_state=42)
#     gb = GradientBoostingClassifier(random_state=42)

#     ensemble = VotingClassifier(
#         estimators=[("lr", lr), ("dt", dt), ("rf", rf), ("gb", gb)],
#         voting="soft"
#     )
#     ensemble.fit(X_train, y_train)
#     return ensemble, le, training_columns


# model, label_encoder, training_columns = train_model()

# # =========================================================
# # OPENAI CLIENT
# # =========================================================
# client = None
# if "OPENAI_API_KEY" in st.secrets:
#     clean_key = st.secrets["OPENAI_API_KEY"]
#     clean_key = clean_key.replace("\u200b", "").strip()
#     clean_key = re.sub(r"[^\x00-\x7F]+", "", clean_key)
#     if clean_key:
#         client = OpenAI(api_key=clean_key)

# # =========================================================
# # SESSION STATE
# # =========================================================
# if "user" not in st.session_state:
#     st.session_state.user = None

# if "page" not in st.session_state:
#     st.session_state.page = "Login"

# # =========================================================
# # MASTER DATA
# # =========================================================
# institutions = [
#     "Abia State University", "Abubakar Tafawa Balewa University",
#     "Adekunle Ajasin University", "Afe Babalola University",
#     "Ahmadu Bello University", "Ajayi Crowther University",
#     "American University of Nigeria", "Babcock University",
#     "Bayero University Kano", "Benson Idahosa University",
#     "Benue State University", "Bowen University", "Covenant University",
#     "Delta State University", "Ekiti State University",
#     "Federal University of Technology Akure",
#     "Federal University of Technology Minna",
#     "Federal University of Technology Owerri",
#     "Igbinedion University", "Kaduna State University",
#     "Kogi State University", "Lagos State University",
#     "Lead City University", "Madonna University",
#     "Nile University of Nigeria", "Nnamdi Azikiwe University",
#     "Obafemi Awolowo University", "Pan Atlantic University",
#     "Rivers State University", "University of Abuja",
#     "University of Benin", "University of Calabar",
#     "University of Ibadan", "University of Ilorin",
#     "University of Jos", "University of Lagos",
#     "University of Maiduguri", "University of Nigeria Nsukka",
#     "University of Port Harcourt", "University of Uyo"
# ]

# programmes = [
#     "Accounting", "Adult Education", "Agricultural Engineering",
#     "Banking and Finance", "Biochemistry", "Biology",
#     "Biomedical Engineering", "Business Administration",
#     "Chemical Engineering", "Chemistry", "Civil Engineering",
#     "Computer Science", "Curriculum Studies", "Cybersecurity",
#     "Data Science", "Economics", "Educational Management",
#     "Electrical Engineering", "Geology", "Guidance and Counselling",
#     "Information Technology", "Mathematics", "Mechanical Engineering",
#     "Mechatronics Engineering", "Microbiology", "Petroleum Engineering",
#     "Physics", "Public Administration", "Software Engineering", "Statistics"
# ]

# disciplines = ["Computing", "Education", "Engineering", "Management", "Science"]

# # =========================================================
# # QUESTION OPTION BANKS
# # =========================================================
# PHILOSOPHY_OPTIONS = [
#     {
#         "label": "Clearly defined and similar to those laid down in the CCMAS for the programme",
#         "score": 1.0,
#     },
#     {
#         "label": "Not well stated/Not in line with those laid down in the CCMAS for the programme",
#         "score": 0.0,
#     },
# ]

# RELATIONSHIP_CCMAS_OPTIONS = [
#     {
#         "label": "The curriculum is adequate for the degree programme as it contains all the core/compulsory courses prescribed by the CCMAS",
#         "score": 1.0,
#     },
#     {
#         "label": "Not fully adequate for the degree programme as a few core/compulsory courses are omitted in some levels of study as prescribed by the CCMAS",
#         "score": 0.5,
#     },
#     {
#         "label": "Not adequate for the degree programme as it does not contain all the core/compulsory courses across all the levels of study as prescribed by the CCMAS",
#         "score": 0.0,
#     },
# ]

# INNOVATION_OPTIONS = [
#     {
#         "label": "Inclusion of four or more additional/innovative courses in the curriculum",
#         "score": 1.0,
#     },
#     {
#         "label": "Inclusion of less than four innovative courses in the curriculum",
#         "score": 0.5,
#     },
#     {
#         "label": "Non-inclusion of additional/innovative courses in the curriculum",
#         "score": 0.0,
#     },
# ]

# COVERAGE_OPTIONS = [
#     {"label": "The curriculum is adequately covered", "score": 1.0},
#     {"label": "The curriculum is substantially covered", "score": 0.5},
#     {"label": "The curriculum is not adequately covered", "score": 0.0},
# ]

# ADMISSION_OPTIONS = [
#     {
#         "label": "All students enrolled in the programme meet the admission requirements",
#         "score": 1.0,
#     },
#     {
#         "label": "Some of the students enrolled in the programme did not meet the admission requirements",
#         "score": 0.0,
#     },
# ]

# ACADEMIC_REG_OPTIONS = [
#     {
#         "label": "Available, quite clear, are in use and well publicized to students",
#         "score": 1.0,
#     },
#     {
#         "label": "Available, not clear, but in use and well publicized to students",
#         "score": 0.5,
#     },
#     {
#         "label": "Not available",
#         "score": 0.0,
#     },
# ]

# TEST_EXAM_OPTIONS = [
#     {
#         "label": "Very good standard and quality and adequately cover the curriculum",
#         "score": 1.0,
#     },
#     {
#         "label": "Good standard and quality but do not cover the curriculum",
#         "score": 0.5,
#     },
#     {
#         "label": "Not of good standard and do not adequately cover the curriculum",
#         "score": 0.0,
#     },
# ]

# EVALUATION_WORK_OPTIONS = [
#     {
#         "label": "Marking schemes exist, are well-developed and the grading of projects, CA, course work and exam scripts is consistent",
#         "score": 1.0,
#     },
#     {
#         "label": "Marking schemes exist, not well-developed and the grading of projects, CA, course work and exam scripts is not consistent",
#         "score": 0.5,
#     },
#     {
#         "label": "Marking schemes do not exist, and the grading of projects, CA, course work and exam script is poor and not consistent",
#         "score": 0.0,
#     },
# ]

# DEGREE_PROJECT_OPTIONS = [
#     {
#         "label": "Good quality, well supervised, innovative topics and creative",
#         "score": 1.0,
#     },
#     {
#         "label": "Innovative and creative but not well supervised",
#         "score": 0.5,
#     },
#     {
#         "label": "Not well supervised, lack creativity and not innovative",
#         "score": 0.0,
#     },
# ]

# PRACTICAL_WORK_OPTIONS = [
#     {
#         "label": "Practicals conducted, good quality, depth and scope covered",
#         "score": 1.0,
#     },
#     {
#         "label": "Practicals conducted, good quality but did not cover the curriculum",
#         "score": 0.66,
#     },
#     {
#         "label": "Practicals conducted but not of good quality",
#         "score": 0.33,
#     },
#     {
#         "label": "Practical not conducted",
#         "score": 0.0,
#     },
# ]

# STUDENTS_COURSE_EVAL_OPTIONS = [
#     {
#         "label": "The course content, learning materials, course delivery, physical facilities are adequate",
#         "score": 1.0,
#     },
#     {
#         "label": "The course content, learning materials, course delivery, physical facilities are fairly adequate",
#         "score": 0.5,
#     },
#     {
#         "label": "The course content, learning materials, course delivery, physical facilities are not adequate",
#         "score": 0.0,
#     },
# ]

# SKILLS_ACQUISITION_OPTIONS = [
#     {
#         "label": "Students demonstrate very good level of hard and soft skills relevant to the programme",
#         "score": 1.0,
#     },
#     {
#         "label": "Students demonstrate good level of hard and soft skills relevant to the programme",
#         "score": 0.66,
#     },
#     {
#         "label": "Students demonstrate average level of hard and soft skills relevant to the programme",
#         "score": 0.33,
#     },
#     {
#         "label": "Students demonstrate poor level of hard and soft skills relevant to the programme",
#         "score": 0.0,
#     },
# ]

# EXTERNAL_EXAM_OPTIONS = [
#     {
#         "label": "External examiners system exists. Qualified assessors are engaged and the work done is of good standard",
#         "score": 1.0,
#     },
#     {
#         "label": "External examiners system exists. Qualified assessors are engaged but the work done is not of good standard",
#         "score": 0.5,
#     },
#     {
#         "label": "External examiners system exists but the quality of assessors engaged is poor OR external examiners system does not exist",
#         "score": 0.0,
#     },
# ]

# IQA_OPTIONS = [
#     {"label": "Exists and effective", "score": 1.0},
#     {"label": "Exists not effective", "score": 0.5},
#     {"label": "Does not exist", "score": 0.0},
# ]

# COMPETENCE_OPTIONS = [
#     {"label": "Competent", "score": 1.0},
#     {"label": "Not competent", "score": 0.0},
# ]

# ADMIN_OPTIONS = [
#     {
#         "label": "Run by a qualified academic staff (Senior Lecturer and above) and very effective and efficient",
#         "score": 1.0,
#     },
#     {
#         "label": "Run by a qualified academic staff (Senior Lecturer and above) and efficient",
#         "score": 0.5,
#     },
#     {
#         "label": "Run by an inexperienced academic and generally ineffective and inefficient",
#         "score": 0.0,
#     },
# ]

# LAB_SPACE_OPTIONS = [
#     {
#         "label": "The space is adequate and meets the provisions of the NUC space standard by 70% or more",
#         "score": 1.0,
#     },
#     {
#         "label": "The space meets 60% but less than 70% of the NUC space standards",
#         "score": 0.66,
#     },
#     {
#         "label": "The space meets 50% but less than 60% of the NUC space standards",
#         "score": 0.33,
#     },
#     {
#         "label": "The space meets less than 50% of the NUC space standards",
#         "score": 0.0,
#     },
# ]

# LAB_EQUIPMENT_OPTIONS = [
#     {"label": "Available equipment are 80% or more", "score": 1.0},
#     {"label": "70% but less than 80%", "score": 0.75},
#     {"label": "60% but less than 70%", "score": 0.5},
#     {"label": "50% but less than 60%", "score": 0.25},
#     {"label": "Less than 50%", "score": 0.0},
# ]

# CLASSROOM_SPACE_OPTIONS = [
#     {
#         "label": "Classroom space available meets the space standards specified in the CCMAS by 70% or more",
#         "score": 1.0,
#     },
#     {"label": "60% but less than 70%", "score": 0.66},
#     {"label": "50% but less than 60%", "score": 0.33},
#     {"label": "Less than 50%", "score": 0.0},
# ]

# CLASSROOM_EQUIPMENT_OPTIONS = [
#     {"label": "Adequate and well maintained", "score": 1.0},
#     {"label": "Adequate but not well maintained/slightly inadequate but well maintained", "score": 0.5},
#     {"label": "Inadequate and not well maintained", "score": 0.0},
# ]

# OFFICE_OPTIONS = [
#     {"label": "Adequate in space and well equipped", "score": 1.0},
#     {"label": "Slightly inadequate in space but well equipped", "score": 0.66},
#     {"label": "Adequate in space but ill-equipped OR inadequate in space but well equipped", "score": 0.33},
#     {"label": "Inadequate in space and ill-equipped or inappropriate", "score": 0.0},
# ]

# SAFETY_OPTIONS = [
#     {
#         "label": "Safe and comply with all Government Laws relating to fire and environmental sanitation including adequate and functional toilet facilities",
#         "score": 1.0,
#     },
#     {
#         "label": "Reasonably safe and comply with most Government Laws relating to fire and environmental sanitation including some functional toilet facilities",
#         "score": 0.5,
#     },
#     {
#         "label": "Unsafe and violate Government Laws relating to fire and environmental sanitation including toilet facilities",
#         "score": 0.0,
#     },
# ]

# LIBRARY_HOLDINGS_OPTIONS = [
#     {
#         "label": "Good quality and very relevant and adequate in number and coverage",
#         "score": 1.0,
#     },
#     {
#         "label": "Good quality and adequate in number and coverage",
#         "score": 0.66,
#     },
#     {
#         "label": "Inadequate in number, quality and coverage",
#         "score": 0.33,
#     },
#     {
#         "label": "Poor quality and inadequate in number and coverage",
#         "score": 0.0,
#     },
# ]

# CURRENCY_OPTIONS = [
#     {"label": "Very current for both recommended text and journals", "score": 1.0},
#     {"label": "Very current for recommended text but current for journals or vice versa", "score": 0.75},
#     {"label": "Current for recommended text and journals", "score": 0.5},
#     {"label": "Current for recommended text but not for journals or vice versa", "score": 0.25},
#     {"label": "Not current at all for both recommended text and journals", "score": 0.0},
# ]

# E_LIBRARY_OPTIONS = [
#     {
#         "label": "Active subscription to at least 2 relevant databases in addition to open source materials",
#         "score": 1.0,
#     },
#     {
#         "label": "Active subscription to at least 1 relevant database in addition to open source materials",
#         "score": 0.75,
#     },
#     {
#         "label": "Access to only open source materials",
#         "score": 0.5,
#     },
#     {
#         "label": "Access to only offline materials",
#         "score": 0.25,
#     },
#     {
#         "label": "No subscription and no access to offline and open source materials",
#         "score": 0.0,
#     },
# ]

# FUNDING_OPTIONS = [
#     {"label": "Very adequate", "score": 1.0},
#     {"label": "Adequate", "score": 0.66},
#     {"label": "Inadequate", "score": 0.33},
#     {"label": "Poor", "score": 0.0},
# ]

# RESEARCH_COLLAB_OPTIONS = [
#     {"label": "Available, multidisciplinary approach engaged", "score": 1.0},
#     {"label": "Available, no multidisciplinary approach, result applied", "score": 0.66},
#     {"label": "Available only", "score": 0.33},
#     {"label": "Not available at all", "score": 0.0},
# ]

# TRACER_OPTIONS = [
#     {
#         "label": "Tracer system is in place and graduates’ performance on the job is very good",
#         "score": 1.0,
#     },
#     {
#         "label": "Tracer system not in place but graduates’ performance on the job is good or vice versa",
#         "score": 0.5,
#     },
#     {
#         "label": "Tracer system not in place and performance below average",
#         "score": 0.0,
#     },
# ]

# TRACER_NO_GRAD_OPTIONS = [
#     {"label": "Very good", "score": 1.0},
#     {"label": "Average", "score": 0.5},
#     {"label": "Below average", "score": 0.0},
# ]

# # =========================================================
# # HELPERS
# # =========================================================
# def ask_scored_question(label: str, options: list, key: str):
#     option_labels = [opt["label"] for opt in options]
#     selected = st.selectbox(
#         label,
#         option_labels,
#         index=None,
#         placeholder="Select an option",
#         key=key,
#     )
#     if selected is None:
#         return None
#     for opt in options:
#         if opt["label"] == selected:
#             return opt["score"]
#     return None


# def pct(numerator: int, denominator: int) -> float:
#     return 0.0 if denominator <= 0 else (numerator / denominator) * 100


# def normalized_from_score(score: int, max_score: int) -> float:
#     if max_score <= 0:
#         return 0.0
#     ratio = score / max_score
#     if ratio >= 0.75:
#         return 1.0
#     if ratio >= 0.4:
#         return 0.5
#     return 0.0


# def calc_score(section: dict):
#     values = list(section.values())
#     if not values or any(v is None for v in values):
#         return None
#     return (sum(values) / len(values)) * 100


# def display_score_card(title: str, value):
#     display_value = "—" if value is None else f"{value:.2f}%"
#     st.markdown(
#         f'<div class="score-card"><h4>{title}</h4><h2>{display_value}</h2></div>',
#         unsafe_allow_html=True
#     )


# def get_feature_importance(_model, _training_columns):
#     try:
#         base_model = _model.named_estimators_["rf"]
#     except Exception:
#         return pd.DataFrame(columns=["feature", "importance"])

#     return pd.DataFrame({
#         "feature": _training_columns,
#         "importance": base_model.feature_importances_
#     }).sort_values(by="importance", ascending=False)


# def inputs_complete(values):
#     return all(v is not None for v in values)

# # =========================================================
# # STAFFING CALCULATION ENGINE
# # =========================================================
# def score_staff_student_ratio(staff_count: int, student_count: int):
#     if staff_count is None or student_count is None or staff_count <= 0:
#         return None, None, None
#     actual_ratio = round(student_count / staff_count)
#     if actual_ratio <= 20:
#         score = 4
#     elif actual_ratio <= 25:
#         score = 3
#     elif actual_ratio <= 30:
#         score = 2
#     elif actual_ratio <= 35:
#         score = 1
#     else:
#         score = 0
#     return score, actual_ratio, normalized_from_score(score, 4)


# def score_core_staff(core_staff_count: int, staff_count: int):
#     if core_staff_count is None or staff_count is None or staff_count <= 0:
#         return None, None, None
#     core_pct = pct(core_staff_count, staff_count)
#     if core_pct >= 75:
#         score = 6
#     elif core_pct >= 60:
#         score = 4
#     elif core_pct >= 50:
#         score = 2
#     else:
#         score = 0
#     return score, core_pct, normalized_from_score(score, 6)


# def score_staff_mix_by_rank(prof_count: int, senior_count: int, lect1_below_count: int):
#     if prof_count is None or senior_count is None or lect1_below_count is None:
#         return None, {"prof_pct": None, "senior_pct": None, "others_pct": None}, None

#     total = prof_count + senior_count + lect1_below_count
#     if total <= 0:
#         return None, {"prof_pct": None, "senior_pct": None, "others_pct": None}, None

#     prof_pct = round(pct(prof_count, total))
#     senior_pct = round(pct(senior_count, total))
#     others_pct = 100 - prof_pct - senior_pct

#     prof_ok = abs(prof_pct - 20) <= 2
#     senior_ok = abs(senior_pct - 35) <= 2
#     others_ok = abs(others_pct - 45) <= 2
#     categories_met = sum([prof_ok, senior_ok, others_ok])

#     if categories_met == 3:
#         score = 5
#     elif categories_met >= 1:
#         score = 3
#     else:
#         score = 0

#     return score, {
#         "prof_pct": prof_pct,
#         "senior_pct": senior_pct,
#         "others_pct": others_pct,
#     }, normalized_from_score(score, 5)


# def score_phd_qualification(phd_count: int, core_staff_count: int):
#     if phd_count is None or core_staff_count is None or core_staff_count <= 0:
#         return None, None, None
#     phd_pct = pct(phd_count, core_staff_count)
#     if phd_pct >= 70:
#         score = 6
#     elif phd_pct >= 60:
#         score = 4
#     elif phd_pct >= 50:
#         score = 2
#     else:
#         score = 0
#     return score, phd_pct, normalized_from_score(score, 6)


# def score_academic_staff_dev(trained_count: int, academic_staff_count: int):
#     if trained_count is None or academic_staff_count is None or academic_staff_count <= 0:
#         return None, None, None
#     dev_pct = pct(trained_count, academic_staff_count)
#     if dev_pct >= 70:
#         score = 5
#     elif dev_pct >= 60:
#         score = 3
#     elif dev_pct >= 50:
#         score = 1
#     else:
#         score = 0
#     return score, dev_pct, normalized_from_score(score, 5)


# def score_non_teaching_staff(status_choice: str):
#     if status_choice is None:
#         return None, None
#     mapping = {
#         "Adequate in number and quality": (3, 1.0),
#         "Not adequate in number but of good quality": (2, 0.5),
#         "Inadequate in number and of poor quality": (0, 0.0),
#     }
#     return mapping[status_choice]


# def score_non_academic_staff_dev(trained_count: int, non_academic_staff_count: int):
#     if trained_count is None or non_academic_staff_count is None or non_academic_staff_count <= 0:
#         return None, None, None
#     dev_pct = pct(trained_count, non_academic_staff_count)
#     if dev_pct >= 70:
#         score = 2
#     elif dev_pct >= 50:
#         score = 1
#     else:
#         score = 0
#     return score, dev_pct, normalized_from_score(score, 2)

# # =========================================================
# # REPORT CLEANER
# # =========================================================
# def clean_report_text(text: str) -> str:
#     banned_phrases = [
#         "Please let me know if you require assistance with specific sections or further elaboration.",
#         "Please let me know if you need assistance with specific sections or further elaboration.",
#         "Please let me know if you need any further assistance.",
#         "Let me know if you need anything else.",
#         "Please let me know if you need anything else.",
#     ]
#     cleaned = text
#     for phrase in banned_phrases:
#         cleaned = cleaned.replace(phrase, "")
#     return cleaned.strip()

# # =========================================================
# # ADVISORY GENERATOR
# # =========================================================
# def generate_advisory_report(
#     institution: str,
#     programme: str,
#     discipline: str,
#     self_study_score: float,
#     predicted_status: str,
#     actual_ratio: int,
#     ratio_score_raw: int,
#     core_pct: float,
#     core_score_raw: int,
#     mix_score_raw: int,
#     phd_pct: float,
#     phd_score_raw: int,
#     acad_dev_pct: float,
#     acad_dev_score_raw: int,
#     non_teach_score_raw: int,
#     non_acad_dev_pct: float,
#     non_acad_dev_score_raw: int,
#     top_weaknesses: str,
# ) -> str:
#     if client is None:
#         return "OpenAI API key is not configured, so the personalized advisory report cannot be generated."

#     prompt = f"""
# You are an accreditation advisory expert.

# Prepare a well-formatted accreditation readiness report using markdown headings and bullet points.

# Institution: {institution}
# Programme: {programme}
# Discipline: {discipline}

# Programme Score: {round(self_study_score, 2)}%
# Predicted Accreditation Status: {predicted_status}

# Calculated staffing indicators:
# - Staff to Student Ratio: 1:{actual_ratio}
# - Staff to Student Ratio Score: {ratio_score_raw}/4
# - Core Staff Percentage: {round(core_pct, 2)}%
# - Core Staff Score: {core_score_raw}/6
# - Staff Mix by Rank Score: {mix_score_raw}/5
# - PhD Holders Percentage: {round(phd_pct, 2)}%
# - Qualifications of Teaching Staff Score: {phd_score_raw}/6
# - Academic Staff Development Percentage: {round(acad_dev_pct, 2)}%
# - Academic Staff Development Score: {acad_dev_score_raw}/5
# - Non-Teaching Staff Score: {non_teach_score_raw}/3
# - Non-Academic Staff Development Percentage: {round(non_acad_dev_pct, 2)}%
# - Non-Academic Staff Development Score: {non_acad_dev_score_raw}/2

# High-impact weaknesses:
# {top_weaknesses}

# Format the report using exactly these sections:

# ## Executive Summary
# ## Key High-Impact Weaknesses
# ## Section-by-Section Improvement Plan
# ## Documentation Checklist
# ## Risk Mitigation Priorities

# Requirements:
# - Be specific and practical
# - Use concise bullet points under each heading
# - Do not add filler text
# - Do not end with any offer of further help
# - Do not include any sentence like “Please let me know if you need anything else”
# """

#     response = client.chat.completions.create(
#         model="gpt-4.1-mini",
#         messages=[
#             {"role": "system", "content": "You are an accreditation advisory expert and must return neatly formatted markdown."},
#             {"role": "user", "content": prompt},
#         ],
#         temperature=0.5,
#     )
#     return clean_report_text(response.choices[0].message.content)

# # =========================================================
# # AUTH PAGES
# # =========================================================
# def render_login():
#     st.subheader("Login")
#     email = st.text_input("Email", key="login_email")
#     password = st.text_input("Password", type="password", key="login_password")

#     c1, c2 = st.columns(2)
#     with c1:
#         if st.button("Login", width="stretch"):
#             user = authenticate_user(email, password)
#             if user:
#                 st.session_state.user = {
#                     "id": user[0],
#                     "full_name": user[1],
#                     "email": user[2],
#                     "role": user[3],
#                 }
#                 st.session_state.page = "New Assessment"
#                 st.rerun()
#             else:
#                 st.error("Invalid email or password.")
#     with c2:
#         if st.button("Go to Signup", width="stretch"):
#             st.session_state.page = "Signup"
#             st.rerun()


# def render_signup():
#     st.subheader("Create Account")
#     full_name = st.text_input("Full Name", key="signup_name")
#     email = st.text_input("Email", key="signup_email")
#     password = st.text_input("Password", type="password", key="signup_password")
#     confirm_password = st.text_input("Confirm Password", type="password", key="signup_confirm")

#     c1, c2 = st.columns(2)
#     with c1:
#         if st.button("Create Account", width="stretch"):
#             if not full_name or not email or not password:
#                 st.error("Please complete all fields.")
#             elif password != confirm_password:
#                 st.error("Passwords do not match.")
#             else:
#                 ok, msg = create_user(full_name, email, password)
#                 if ok:
#                     st.success(msg)
#                     st.session_state.page = "Login"
#                     st.rerun()
#                 else:
#                     st.error(msg)
#     with c2:
#         if st.button("Back to Login", width="stretch"):
#             st.session_state.page = "Login"
#             st.rerun()

# # =========================================================
# # PRE-LOGIN
# # =========================================================
# if st.session_state.user is None:
#     auth_tab1, auth_tab2 = st.tabs(["Login", "Signup"])
#     with auth_tab1:
#         render_login()
#     with auth_tab2:
#         render_signup()
#     st.stop()

# # =========================================================
# # SIDEBAR NAV
# # =========================================================
# with st.sidebar:
#     st.markdown(f"### Welcome, {st.session_state.user['full_name']}")
#     st.caption(st.session_state.user["email"])

#     nav_options = ["New Assessment", "Saved Assessments", "Report History"]
#     if st.session_state.user.get("role") == "admin":
#         nav_options = ["New Assessment", "Saved Assessments", "Report History", "Admin Dashboard"]

#     choice = st.radio("Navigation", nav_options)
#     st.session_state.page = choice

#     if st.button("Logout", width="stretch"):
#         st.session_state.user = None
#         st.session_state.page = "Login"
#         st.rerun()

# # =========================================================
# # NEW ASSESSMENT
# # =========================================================
# if st.session_state.page == "New Assessment":
#     st.header("New Assessment")

#     c1, c2, c3 = st.columns(3)
#     with c1:
#         institution = st.selectbox("Name of Institution", institutions, index=None, placeholder="Select institution")
#     with c2:
#         discipline = st.selectbox("Discipline", disciplines, index=None, placeholder="Select discipline")
#     with c3:
#         programme = st.selectbox("Programme", programmes, index=None, placeholder="Select programme")

#     st.subheader("Staffing & Enrollment Data")

#     s1, s2, s3 = st.columns(3)
#     with s1:
#         academic_staff_count = st.number_input("Number of academic staff", min_value=1, value=None, placeholder="Enter value")
#     with s2:
#         student_count = st.number_input("Number of students", min_value=1, value=None, placeholder="Enter value")
#     with s3:
#         core_staff_count = st.number_input("Number of academic staff core to the subject area", min_value=0, value=None, placeholder="Enter value")

#     s4, s5, s6 = st.columns(3)
#     with s4:
#         professor_reader_count = st.number_input("Number of Professors/Readers", min_value=0, value=None, placeholder="Enter value")
#     with s5:
#         senior_lecturer_count = st.number_input("Number of Senior Lecturers", min_value=0, value=None, placeholder="Enter value")
#     with s6:
#         lecturer1_below_count = st.number_input("Number of Lecturers I and below", min_value=0, value=None, placeholder="Enter value")

#     s7, s8, s9 = st.columns(3)
#     with s7:
#         phd_holder_count = st.number_input("Number of Ph.D Holders", min_value=0, value=None, placeholder="Enter value")
#     with s8:
#         academic_staff_dev_count = st.number_input("Number of academic staff with staff development programme", min_value=0, value=None, placeholder="Enter value")
#     with s9:
#         non_academic_staff_count = st.number_input("Number of non-teaching staff", min_value=0, value=None, placeholder="Enter value")

#     non_academic_staff_dev_count = st.number_input(
#         "Number of non-academic staff with staff development programme",
#         min_value=0,
#         value=None,
#         placeholder="Enter value"
#     )

#     non_teaching_quality_choice = st.selectbox(
#         "Non-Teaching Staff",
#         [
#             "Adequate in number and quality",
#             "Not adequate in number but of good quality",
#             "Inadequate in number and of poor quality",
#         ],
#         index=None,
#         placeholder="Select status"
#     )

#     staffing_inputs_complete = inputs_complete([
#         academic_staff_count, student_count, core_staff_count,
#         professor_reader_count, senior_lecturer_count, lecturer1_below_count,
#         phd_holder_count, academic_staff_dev_count,
#         non_academic_staff_count, non_academic_staff_dev_count,
#         non_teaching_quality_choice
#     ])

#     if staffing_inputs_complete:
#         ratio_score_raw, actual_ratio, staff_ratio_feature = score_staff_student_ratio(
#             academic_staff_count, student_count
#         )
#         core_score_raw, core_pct, core_staff_feature = score_core_staff(
#             core_staff_count, academic_staff_count
#         )
#         mix_score_raw, mix_pct_dict, staff_mix_feature = score_staff_mix_by_rank(
#             professor_reader_count, senior_lecturer_count, lecturer1_below_count
#         )
#         phd_score_raw, phd_pct, phd_feature = score_phd_qualification(
#             phd_holder_count, core_staff_count
#         )
#         acad_dev_score_raw, acad_dev_pct, acad_dev_feature = score_academic_staff_dev(
#             academic_staff_dev_count, academic_staff_count
#         )
#         non_teach_score_raw, non_teach_feature = score_non_teaching_staff(non_teaching_quality_choice)
#         non_acad_dev_score_raw, non_acad_dev_pct, non_acad_dev_feature = score_non_academic_staff_dev(
#             non_academic_staff_dev_count, non_academic_staff_count
#         )

#         with st.expander("Computed Staffing Indicators", expanded=True):
#             m1, m2, m3 = st.columns(3)
#             with m1:
#                 st.metric("Staff to Student Ratio", f"1 : {actual_ratio}")
#                 st.metric("Staff to Student Ratio Score", f"{ratio_score_raw}/4")
#             with m2:
#                 st.metric("Percentage of Staff Core to the Subject Area", f"{round(core_pct, 2)}%")
#                 st.metric("Proportion of Staff Core to the Subject Area Score", f"{core_score_raw}/6")
#             with m3:
#                 st.metric("Ph.D Holders Percentage", f"{round(phd_pct, 2)}%")
#                 st.metric("Qualifications of Teaching Staff Score", f"{phd_score_raw}/6")

#             m4, m5, m6 = st.columns(3)
#             with m4:
#                 st.metric(
#                     "Staff Mix by Rank",
#                     f"{mix_pct_dict['prof_pct']}:{mix_pct_dict['senior_pct']}:{mix_pct_dict['others_pct']}"
#                 )
#                 st.metric("Staff Mix by Rank Score", f"{mix_score_raw}/5")
#             with m5:
#                 st.metric("Academic Staff Development Percentage", f"{round(acad_dev_pct, 2)}%")
#                 st.metric("Academic Staff Development Score", f"{acad_dev_score_raw}/5")
#             with m6:
#                 st.metric("Non-Academic Staff Development Percentage", f"{round(non_acad_dev_pct, 2)}%")
#                 st.metric("Non-Academic Staff Development Score", f"{non_acad_dev_score_raw}/2")

#             st.metric("Non-Teaching Staff Score", f"{non_teach_score_raw}/3")
#     else:
#         ratio_score_raw = actual_ratio = staff_ratio_feature = None
#         core_score_raw = core_pct = core_staff_feature = None
#         mix_score_raw = staff_mix_feature = None
#         mix_pct_dict = {"prof_pct": None, "senior_pct": None, "others_pct": None}
#         phd_score_raw = phd_pct = phd_feature = None
#         acad_dev_score_raw = acad_dev_pct = acad_dev_feature = None
#         non_teach_score_raw = non_teach_feature = None
#         non_acad_dev_score_raw = non_acad_dev_pct = non_acad_dev_feature = None
#         st.info("Enter all staffing and enrollment figures to compute staffing indicators.")

#     tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
#         "Academic Content",
#         "Staffing",
#         "Physical Facilities",
#         "Library",
#         "Funding",
#         "Research and Collaboration",
#         "Tracer and Employers’ Rating",
#     ])

#     with tab1:
#         academic = {
#             "philosophy_objectives_defined": ask_scored_question(
#                 "Philosophy and Objectives of the Programme",
#                 PHILOSOPHY_OPTIONS,
#                 "a0"
#             ),
#             "curriculum_aligned_with_BMAS": ask_scored_question(
#                 "Relationship between CCMAS and Curriculum",
#                 RELATIONSHIP_CCMAS_OPTIONS,
#                 "a1"
#             ),
#             "innovative_courses_present": ask_scored_question(
#                 "Innovation (Additional Courses)",
#                 INNOVATION_OPTIONS,
#                 "a2"
#             ),
#             "curriculum_coverage_complete": ask_scored_question(
#                 "Coverage of the Curriculum",
#                 COVERAGE_OPTIONS,
#                 "a3"
#             ),
#             "admission_requirements_compliant": ask_scored_question(
#                 "Admission Requirements into the Programme",
#                 ADMISSION_OPTIONS,
#                 "a4"
#             ),
#             "academic_regulations_defined": ask_scored_question(
#                 "Academic Regulations",
#                 ACADEMIC_REG_OPTIONS,
#                 "a5"
#             ),
#             "tests_and_examinations_standardized": ask_scored_question(
#                 "Standard of Tests and Examinations",
#                 TEST_EXAM_OPTIONS,
#                 "a6"
#             ),
#             "evaluation_methods_clear": ask_scored_question(
#                 "Evaluation of Students’ Work",
#                 EVALUATION_WORK_OPTIONS,
#                 "a7"
#             ),
#             "degree_projects_adequate": ask_scored_question(
#                 "Degree Projects",
#                 DEGREE_PROJECT_OPTIONS,
#                 "a8"
#             ),
#             "practical_work_adequate": ask_scored_question(
#                 "Practical Work",
#                 PRACTICAL_WORK_OPTIONS,
#                 "a9"
#             ),
#             "student_course_evaluation_present": ask_scored_question(
#                 "Students’ Course Evaluation",
#                 STUDENTS_COURSE_EVAL_OPTIONS,
#                 "a10"
#             ),
#             "skills_acquisition_programme": ask_scored_question(
#                 "Evaluation of Skills Acquisition",
#                 SKILLS_ACQUISITION_OPTIONS,
#                 "a11"
#             ),
#             "external_examiner_system": ask_scored_question(
#                 "External Examination System",
#                 EXTERNAL_EXAM_OPTIONS,
#                 "a12"
#             ),
#             "internal_quality_assurance": ask_scored_question(
#                 "Internal Quality Assurance System",
#                 IQA_OPTIONS,
#                 "a13"
#             ),
#         }

#     with tab2:
#         staffing = {
#             "proportion_core_staff_sufficient": core_staff_feature,
#             "staff_rank_mix_balanced": staff_mix_feature,
#             "academic_staff_qualification_high": phd_feature,
#             "staff_competence_verified": ask_scored_question(
#                 "Competence of Teaching Staff",
#                 COMPETENCE_OPTIONS,
#                 "s1"
#             ),
#             "administrative_support_available": ask_scored_question(
#                 "Administration of College/School/Faculty/Department",
#                 ADMIN_OPTIONS,
#                 "s2"
#             ),
#             "non_teaching_staff_adequate": non_teach_feature,
#             "academic_staff_development_programme": acad_dev_feature,
#             "non_academic_staff_development_programme": non_acad_dev_feature,
#         }
#         st.info("Calculated staffing items are automatically derived from the figures entered above.")

#     with tab3:
#         facilities = {
#             "laboratory_space_adequate": ask_scored_question(
#                 "Laboratories/Clinics/Studios/Farms/Museums",
#                 LAB_SPACE_OPTIONS,
#                 "f1"
#             ),
#             "laboratory_equipment_adequate": ask_scored_question(
#                 "Laboratory Equipment",
#                 LAB_EQUIPMENT_OPTIONS,
#                 "f2"
#             ),
#             "classroom_space_adequate": ask_scored_question(
#                 "Classrooms/Lecture Theatres",
#                 CLASSROOM_SPACE_OPTIONS,
#                 "f3"
#             ),
#             "classroom_equipment_adequate": ask_scored_question(
#                 "Classroom Equipment",
#                 CLASSROOM_EQUIPMENT_OPTIONS,
#                 "f4"
#             ),
#             "office_accommodation_adequate": ask_scored_question(
#                 "Office Accommodation",
#                 OFFICE_OPTIONS,
#                 "f5"
#             ),
#             "safety_environment_present": ask_scored_question(
#                 "Safety and Environment",
#                 SAFETY_OPTIONS,
#                 "f6"
#             ),
#         }

#     with tab4:
#         library = {
#             "library_holdings_adequate": ask_scored_question(
#                 "Holdings",
#                 LIBRARY_HOLDINGS_OPTIONS,
#                 "l1"
#             ),
#             "library_material_current": ask_scored_question(
#                 "Currency of Holdings",
#                 CURRENCY_OPTIONS,
#                 "l2"
#             ),
#             "e_library_subscription_available": ask_scored_question(
#                 "Subscription to e-Books and e-Journals",
#                 E_LIBRARY_OPTIONS,
#                 "l3"
#             ),
#             "e_library_access_good": ask_scored_question(
#                 "Access to available e-Books and e-Journals",
#                 E_LIBRARY_OPTIONS,
#                 "l4"
#             ),
#         }

#     with tab5:
#         funding = {
#             "programme_funding_adequate": ask_scored_question(
#                 "Funding",
#                 FUNDING_OPTIONS,
#                 "fd0"
#             ),
#             "budget_release_regular": 1.0,
#             "equipment_maintenance_budget_available": 1.0,
#         }

#     with tab6:
#         research_collaboration = {
#             "research_collaboration_active": ask_scored_question(
#                 "Research and Collaboration",
#                 RESEARCH_COLLAB_OPTIONS,
#                 "r1"
#             ),
#             "research_output_present": 1.0,
#         }

#     with tab7:
#         has_graduated_students = st.selectbox(
#             "Has the programme graduated students?",
#             ["Yes", "No"],
#             index=None,
#             placeholder="Select an option",
#             key="graduated_students"
#         )

#         tracer_options = None
#         if has_graduated_students == "Yes":
#             tracer_options = TRACER_OPTIONS
#         elif has_graduated_students == "No":
#             tracer_options = TRACER_NO_GRAD_OPTIONS

#         tracer_rating = {
#             "employer_rating_positive": ask_scored_question(
#                 "Tracer and Employers’ Rating",
#                 tracer_options,
#                 "t1"
#             ) if tracer_options else None,
#             "tracer_study_available": 1.0 if has_graduated_students == "Yes" else 0.5 if has_graduated_students == "No" else None,
#         }

#     academic_score = calc_score(academic)

#     staffing_display_items = {
#         "staff_student_ratio_compliant": staff_ratio_feature,
#         "proportion_core_staff_sufficient": core_staff_feature,
#         "staff_rank_mix_balanced": staff_mix_feature,
#         "academic_staff_qualification_high": phd_feature,
#         "staff_competence_verified": staffing["staff_competence_verified"],
#         "administrative_support_available": staffing["administrative_support_available"],
#         "non_teaching_staff_adequate": non_teach_feature,
#         "academic_staff_development_programme": acad_dev_feature,
#         "non_academic_staff_development_programme": non_acad_dev_feature,
#     }

#     staffing_score = calc_score(staffing_display_items)
#     physical_facilities_score = calc_score(facilities)
#     library_score = calc_score(library)
#     funding_score = calc_score(funding)
#     research_collaboration_score = calc_score(research_collaboration)
#     tracer_employers_score = calc_score(tracer_rating)

#     if all(v is not None for v in [
#         academic_score, staffing_score, physical_facilities_score,
#         library_score, funding_score, research_collaboration_score,
#         tracer_employers_score
#     ]):
#         self_study_score = (
#             academic_score * 0.25 +
#             staffing_score * 0.25 +
#             physical_facilities_score * 0.15 +
#             library_score * 0.10 +
#             funding_score * 0.10 +
#             research_collaboration_score * 0.10 +
#             tracer_employers_score * 0.05
#         )
#     else:
#         self_study_score = None

#     st.subheader("Predict Accreditation Status")

#     required_basic = inputs_complete([institution, discipline, programme])
#     required_all_sections = all(v is not None for v in [
#         academic_score, staffing_score, physical_facilities_score,
#         library_score, funding_score, research_collaboration_score,
#         tracer_employers_score, self_study_score
#     ])

#     if st.button("Predict Accreditation Status", type="primary", width="stretch"):
#         if not required_basic:
#             st.error("Please complete the institutional information section.")
#             st.stop()

#         if not staffing_inputs_complete:
#             st.error("Please complete all staffing and enrollment inputs.")
#             st.stop()

#         if not required_all_sections:
#             st.error("Please complete all assessment sections before generating the result.")
#             st.stop()

#         input_data = {}
#         input_data.update(academic)
#         input_data.update(staffing)
#         input_data.update(facilities)
#         input_data.update(library)
#         input_data.update(funding)
#         input_data.update(research_collaboration)
#         input_data.update(tracer_rating)
#         input_data["staff_student_ratio_compliant"] = staff_ratio_feature
#         input_data["discipline"] = discipline

#         input_df = pd.DataFrame([input_data])
#         input_encoded = pd.get_dummies(input_df)
#         input_encoded = input_encoded.reindex(columns=training_columns, fill_value=0)

#         prediction = model.predict(input_encoded)
#         predicted_status = label_encoder.inverse_transform(prediction)[0]

#         st.subheader("Section Scores")

#         row1 = st.columns(4)
#         row2 = st.columns(4)

#         with row1[0]:
#             display_score_card("Academic Content", academic_score)
#         with row1[1]:
#             display_score_card("Staffing", staffing_score)
#         with row1[2]:
#             display_score_card("Physical Facilities", physical_facilities_score)
#         with row1[3]:
#             display_score_card("Library", library_score)

#         with row2[0]:
#             display_score_card("Funding", funding_score)
#         with row2[1]:
#             display_score_card("Research and Collaboration", research_collaboration_score)
#         with row2[2]:
#             display_score_card("Tracer and Employers’ Rating", tracer_employers_score)
#         with row2[3]:
#             display_score_card("Programme Score", self_study_score)

#         st.subheader("Predicted Accreditation Outcome")
#         st.success(predicted_status)

#         importance_df = get_feature_importance(model, training_columns)
#         weak = []
#         for col in input_df.columns:
#             if col in importance_df["feature"].values:
#                 value = input_df[col].iloc[0]
#                 if value in [0.0, 0.5]:
#                     impact = importance_df.loc[
#                         importance_df["feature"] == col,
#                         "importance"
#                     ].values[0]
#                     weak.append((col, value, impact))

#         weak = sorted(weak, key=lambda x: x[2], reverse=True)

#         if weak:
#             st.subheader("Top High-Impact Weak Areas")
#             weak_df = pd.DataFrame([
#                 {
#                     "Feature": w[0].replace("_", " ").title(),
#                     "Severity": "Critical" if w[1] == 0.0 else "Moderate",
#                     "Importance": round(float(w[2]), 4),
#                 }
#                 for w in weak[:10]
#             ])
#             st.dataframe(weak_df, width="stretch")
#             top_weaknesses = "\n".join(
#                 [f"- {w[0].replace('_', ' ').title()}" for w in weak[:10]]
#             )
#         else:
#             top_weaknesses = "- No major weaknesses were detected from the submitted responses."
#             st.info("No major weaknesses were detected from the submitted responses.")

#         with st.spinner("Generating personalized advisory report..."):
#             report = generate_advisory_report(
#                 institution=institution,
#                 programme=programme,
#                 discipline=discipline,
#                 self_study_score=self_study_score,
#                 predicted_status=predicted_status,
#                 actual_ratio=actual_ratio,
#                 ratio_score_raw=ratio_score_raw,
#                 core_pct=core_pct,
#                 core_score_raw=core_score_raw,
#                 mix_score_raw=mix_score_raw,
#                 phd_pct=phd_pct,
#                 phd_score_raw=phd_score_raw,
#                 acad_dev_pct=acad_dev_pct,
#                 acad_dev_score_raw=acad_dev_score_raw,
#                 non_teach_score_raw=non_teach_score_raw,
#                 non_acad_dev_pct=non_acad_dev_pct,
#                 non_acad_dev_score_raw=non_acad_dev_score_raw,
#                 top_weaknesses=top_weaknesses,
#             )

#         st.subheader("Personalized Accreditation Advisory Report")
#         st.markdown(f'<div class="report-box">{report}</div>', unsafe_allow_html=True)

#         save_assessment(
#             user_id=st.session_state.user["id"],
#             institution=institution,
#             discipline=discipline,
#             programme=programme,
#             self_study_score=float(self_study_score),
#             predicted_status=predicted_status,
#             actual_ratio=float(actual_ratio),
#             core_pct=float(core_pct),
#             phd_pct=float(phd_pct),
#             assessment_payload=json.dumps(input_data),
#             report_text=report,
#         )

#         st.success("Assessment and report saved successfully.")

# # =========================================================
# # SAVED ASSESSMENTS
# # =========================================================
# elif st.session_state.page == "Saved Assessments":
#     st.header("Saved Assessments")
#     df_assessments = get_user_assessments(st.session_state.user["id"])

#     if df_assessments.empty:
#         st.info("No saved assessments yet.")
#     else:
#         st.dataframe(df_assessments, width="stretch")

# # =========================================================
# # REPORT HISTORY
# # =========================================================
# elif st.session_state.page == "Report History":
#     st.header("Report History")
#     df_reports = get_user_report_history(st.session_state.user["id"])

#     if df_reports.empty:
#         st.info("No report history yet.")
#     else:
#         for _, row in df_reports.iterrows():
#             title = f"{row['institution_name']} | {row['programme_name']} | {row['predicted_status']} | {row['created_at']}"
#             with st.expander(title):
#                 st.write(f"**Programme Score:** {row['self_study_score']:.2f}%")
#                 st.markdown(f'<div class="report-box">{row["report_text"]}</div>', unsafe_allow_html=True)

# # =========================================================
# # ADMIN DASHBOARD
# # =========================================================
# elif st.session_state.page == "Admin Dashboard" and st.session_state.user.get("role") == "admin":
#     st.header("Admin Dashboard")

#     users_df = get_all_users()
#     assessments_df = get_all_assessments()

#     d1, d2, d3 = st.columns(3)
#     with d1:
#         st.metric("Total Users", len(users_df))
#     with d2:
#         st.metric("Total Assessments", len(assessments_df))
#     with d3:
#         if not assessments_df.empty:
#             st.metric("Average Programme Score", f"{assessments_df['self_study_score'].mean():.2f}%")
#         else:
#             st.metric("Average Programme Score", "0.00%")

#     st.subheader("Users")
#     st.dataframe(users_df, width="stretch")

#     st.subheader("All Assessments")
#     st.dataframe(assessments_df, width="stretch")



import re
import json
import sqlite3
import hashlib
import hmac
from datetime import datetime
from dataclasses import dataclass, asdict
from html import escape
from io import BytesIO
from pathlib import Path

import streamlit as st
import pandas as pd
from openai import OpenAI

from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder
from sklearn.ensemble import RandomForestClassifier

from config.computing_requirements import COMPUTING_PROGRAMMES, COMPUTING_RULES, DOCUMENT_CATEGORIES
from config.institutions import KNOWN_INSTITUTIONS
from services.document_parser import read_uploaded_file
from services.institution_extractor import extract_institution_name, institution_is_resolved

try:
    import bcrypt
except ImportError:
    bcrypt = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

# =========================================================
# PAGE CONFIG
# =========================================================
st.set_page_config(
    page_title="NUC Accreditation Prediction System",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)

# =========================================================
# CUSTOM STYLING
# =========================================================
st.markdown("""
<style>
    :root {
        --nuc-navy: #0b1f3a;
        --nuc-blue: #123a63;
        --nuc-green: #0f6b4f;
        --nuc-gold: #c9a227;
        --nuc-grey: #f4f7f9;
        --nuc-border: #d9e2ea;
        --nuc-text: #102033;
    }

    .stApp {
        background: linear-gradient(180deg, #f6f9fc 0%, #eef3f6 100%);
        color: var(--nuc-text);
    }

    .block-container {
        padding-top: 1.4rem;
        padding-bottom: 2rem;
        max-width: 1220px;
    }

    .hero-card {
        padding: 1.35rem 1.45rem;
        border-radius: 8px;
        border: 1px solid rgba(11,31,58,0.12);
        background: linear-gradient(135deg, rgba(11,31,58,0.96), rgba(15,107,79,0.88));
        color: #ffffff;
        margin-bottom: 1.1rem;
        box-shadow: 0 12px 28px rgba(11,31,58,0.14);
    }

    .hero-card h1,
    .hero-card h2,
    .hero-card h3,
    .hero-card p {
        color: #ffffff;
    }

    .institution-card {
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid var(--nuc-border);
        background: rgba(255,255,255,0.92);
        box-shadow: 0 8px 22px rgba(16,32,51,0.07);
        margin-bottom: 1rem;
    }

    .auth-panel {
        padding: 1.45rem;
        border-radius: 8px;
        border: 1px solid rgba(217,226,234,0.95);
        background: #ffffff;
        box-shadow: 0 16px 36px rgba(11,31,58,0.12);
    }

    .benefit-item {
        display: flex;
        gap: 0.55rem;
        align-items: center;
        margin: 0.55rem 0;
        color: #26384f;
        font-weight: 600;
    }

    .benefit-dot {
        width: 10px;
        height: 10px;
        border-radius: 999px;
        background: var(--nuc-gold);
        flex: 0 0 auto;
    }

    .section-title {
        margin: 0.2rem 0 0.35rem;
        color: var(--nuc-navy);
        font-weight: 750;
    }

    .subtle-copy {
        color: #53657a;
        line-height: 1.55;
    }

    .sidebar-brand {
        padding: 0.75rem 0 1rem;
        border-bottom: 1px solid rgba(217,226,234,0.75);
        margin-bottom: 0.75rem;
    }

    .user-chip {
        padding: 0.75rem;
        border-radius: 8px;
        border: 1px solid rgba(217,226,234,0.9);
        background: rgba(255,255,255,0.76);
        margin: 0.75rem 0 1rem;
    }

    .avatar {
        width: 42px;
        height: 42px;
        display: inline-flex;
        align-items: center;
        justify-content: center;
        border-radius: 50%;
        background: var(--nuc-navy);
        color: white;
        font-weight: 800;
        margin-right: 0.55rem;
    }

    .role-badge {
        display: inline-block;
        padding: 0.16rem 0.5rem;
        border-radius: 999px;
        background: rgba(15,107,79,0.12);
        color: var(--nuc-green);
        border: 1px solid rgba(15,107,79,0.22);
        font-size: 0.78rem;
        font-weight: 750;
        text-transform: uppercase;
    }

    .status-pill {
        display: inline-block;
        padding: 0.2rem 0.62rem;
        border-radius: 999px;
        background: rgba(201,162,39,0.14);
        color: #725a0e;
        border: 1px solid rgba(201,162,39,0.28);
        font-weight: 750;
    }

    .score-card {
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid rgba(15,107,79,0.16);
        background: linear-gradient(180deg, #ffffff, #f3faf7);
        text-align: center;
        min-height: 120px;
    }

    .report-box {
        padding: 1.2rem;
        border-radius: 8px;
        border: 1px solid rgba(11,31,58,0.12);
        background: #ffffff;
        line-height: 1.7;
    }

    div[data-testid="stMetric"] {
        background: #ffffff;
        border: 1px solid rgba(217,226,234,0.92);
        padding: 0.85rem;
        border-radius: 8px;
        box-shadow: 0 5px 16px rgba(16,32,51,0.06);
    }

    .stButton > button {
        border-radius: 8px;
        padding: 0.65rem 1rem;
        font-weight: 700;
        border-color: rgba(11,31,58,0.18);
    }

    .stButton > button[kind="primary"] {
        background: var(--nuc-green);
        border-color: var(--nuc-green);
    }

    section[data-testid="stSidebar"] {
        background: linear-gradient(180deg, #ffffff, #f2f6f8);
        border-right: 1px solid var(--nuc-border);
    }

    section[data-testid="stSidebar"] [role="radiogroup"] label {
        border-radius: 8px;
        padding: 0.2rem 0.35rem;
        margin: 0.08rem 0;
    }

    div[data-testid="stDataFrame"] {
        border: 1px solid var(--nuc-border);
        border-radius: 8px;
        overflow: hidden;
    }

    .stTabs [data-baseweb="tab-list"] {
        gap: 0.35rem;
    }

    .stTabs [data-baseweb="tab"] {
        border-radius: 8px 8px 0 0;
        background: rgba(255,255,255,0.7);
        border: 1px solid rgba(217,226,234,0.75);
    }
</style>
""", unsafe_allow_html=True)

# =========================================================
# DATABASE
# =========================================================
DB_PATH = "nuc_app.db"


def get_conn():
    return sqlite3.connect(DB_PATH, check_same_thread=False)


def init_db():
    conn = get_conn()
    cur = conn.cursor()

    try:
        cur.execute("BEGIN")
        cur.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                full_name TEXT NOT NULL,
                email TEXT NOT NULL UNIQUE,
                password_hash TEXT NOT NULL,
                role TEXT NOT NULL DEFAULT 'user',
                created_at TEXT NOT NULL
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS assessments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                institution_name TEXT NOT NULL,
                discipline TEXT NOT NULL,
                programme_name TEXT NOT NULL,
                programme_score REAL,
                predicted_status TEXT NOT NULL,
                actual_ratio REAL,
                core_pct REAL,
                phd_pct REAL,
                created_at TEXT NOT NULL,
                assessment_payload TEXT NOT NULL,
                report_text TEXT,
                FOREIGN KEY(user_id) REFERENCES users(id)
            )
        """)

        cur.execute("PRAGMA table_info(assessments)")
        existing_cols = {row[1] for row in cur.fetchall()}
        migrations = {
            "programme_score": "REAL",
            "assessment_type": "TEXT DEFAULT 'standard_manual'",
            "rule_based_status": "TEXT",
            "model_name": "TEXT DEFAULT 'RandomForestClassifier'",
            "model_version": "TEXT DEFAULT 'rf-2026-07-13'",
            "ml_confidence": "REAL",
            "extraction_summary": "TEXT",
            "institution_evidence": "TEXT",
            "extracted_institution_value": "TEXT",
            "confirmed_institution_value": "TEXT",
            "institution_corrected_by_user": "INTEGER DEFAULT 0",
            "institution_correction_timestamp": "TEXT",
            "evidence_coverage": "REAL",
            "total_points": "REAL",
            "updated_at": "TEXT",
        }
        for column, ddl in migrations.items():
            if column not in existing_cols:
                cur.execute(f"ALTER TABLE assessments ADD COLUMN {column} {ddl}")

        cur.execute("PRAGMA table_info(assessments)")
        refreshed_cols = {row[1] for row in cur.fetchall()}
        if "self_study_score" in refreshed_cols and "programme_score" in refreshed_cols:
            cur.execute("""
                UPDATE assessments
                SET programme_score = self_study_score
                WHERE programme_score IS NULL
            """)
        cur.execute("""
            UPDATE assessments
            SET rule_based_status = COALESCE(rule_based_status, predicted_status),
                assessment_type = COALESCE(assessment_type, 'standard_manual'),
                updated_at = COALESCE(updated_at, created_at)
        """)
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


init_db()


def hash_password(password: str) -> str:
    if bcrypt is not None:
        return bcrypt.hashpw(password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")
    return hashlib.sha256(password.encode("utf-8")).hexdigest()


def legacy_hash_password(password: str) -> str:
    return hashlib.sha256(password.encode("utf-8")).hexdigest()


def verify_password(password: str, stored_hash: str) -> bool:
    if stored_hash.startswith("$2") and bcrypt is not None:
        return bcrypt.checkpw(password.encode("utf-8"), stored_hash.encode("utf-8"))
    return hmac.compare_digest(stored_hash, legacy_hash_password(password))


def upgrade_password_hash(user_id: int, password: str):
    if bcrypt is None:
        return
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("UPDATE users SET password_hash = ? WHERE id = ?", (hash_password(password), user_id))
    conn.commit()
    conn.close()


def create_user(full_name: str, email: str, password: str):
    conn = get_conn()
    cur = conn.cursor()
    try:
        cur.execute(
            """
            INSERT INTO users (full_name, email, password_hash, role, created_at)
            VALUES (?, ?, ?, ?, ?)
            """,
            (
                full_name,
                email.strip().lower(),
                hash_password(password),
                "user",
                datetime.utcnow().isoformat()
            )
        )
        conn.commit()
        return True, "Account creation is successful."
    except sqlite3.IntegrityError:
        return False, "An account with that email already exists."
    finally:
        conn.close()


def authenticate_user(email: str, password: str):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        """
        SELECT id, full_name, email, role, password_hash
        FROM users
        WHERE email = ?
        """,
        (email.strip().lower(),)
    )
    row = cur.fetchone()
    conn.close()
    if not row or not verify_password(password, row[4]):
        return None
    if not row[4].startswith("$2"):
        upgrade_password_hash(row[0], password)
    return row[:4]


def save_assessment(user_id: int, institution: str, discipline: str, programme: str,
                    programme_score: float, predicted_status: str, actual_ratio: float,
                    core_pct: float, phd_pct: float, assessment_payload: str, report_text: str,
                    assessment_type: str = "standard_manual", rule_based_status: str = None,
                    ml_confidence: float = None, extraction_summary: str = None,
                    evidence_coverage: float = None, total_points: float = None,
                    institution_evidence: str = None, extracted_institution_value: str = None,
                    confirmed_institution_value: str = None,
                    institution_corrected_by_user: bool = False,
                    institution_correction_timestamp: str = None):
    conn = get_conn()
    cur = conn.cursor()
    created_at = datetime.utcnow().isoformat()
    columns = [
        "user_id", "institution_name", "discipline", "programme_name",
        "programme_score", "predicted_status", "actual_ratio", "core_pct", "phd_pct",
        "created_at", "assessment_payload", "report_text", "assessment_type",
        "rule_based_status", "model_name", "model_version", "ml_confidence",
        "extraction_summary", "institution_evidence", "extracted_institution_value",
        "confirmed_institution_value", "institution_corrected_by_user",
        "institution_correction_timestamp", "evidence_coverage", "total_points", "updated_at",
    ]
    values = [
        user_id, institution, discipline, programme,
        programme_score, predicted_status, actual_ratio, core_pct, phd_pct,
        created_at, assessment_payload, report_text, assessment_type,
        rule_based_status or predicted_status,
        "RandomForestClassifier", "rf-2026-07-13", ml_confidence,
        extraction_summary, institution_evidence, extracted_institution_value,
        confirmed_institution_value or institution,
        int(bool(institution_corrected_by_user)), institution_correction_timestamp,
        evidence_coverage, total_points, created_at,
    ]

    cur.execute("PRAGMA table_info(assessments)")
    existing_cols = {row[1] for row in cur.fetchall()}
    if "self_study_score" in existing_cols and "self_study_score" not in columns:
        columns.insert(5, "self_study_score")
        values.insert(5, programme_score)

    placeholders = ", ".join("?" for _ in columns)
    cur.execute(
        f"INSERT INTO assessments ({', '.join(columns)}) VALUES ({placeholders})",
        values,
    )
    conn.commit()
    conn.close()


def get_user_assessments(user_id: int):
    conn = get_conn()
    df = pd.read_sql_query(
        """
        SELECT id, institution_name, discipline, programme_name,
               assessment_type, total_points, programme_score,
               rule_based_status, predicted_status, ml_confidence,
               evidence_coverage, actual_ratio, core_pct, phd_pct, created_at
        FROM assessments
        WHERE user_id = ?
        ORDER BY datetime(created_at) DESC
        """,
        conn,
        params=(user_id,)
    )
    conn.close()
    return df


def get_user_report_history(user_id: int):
    conn = get_conn()
    df = pd.read_sql_query(
        """
        SELECT id, institution_name, discipline, programme_name,
               assessment_type, rule_based_status, predicted_status,
               programme_score, ml_confidence, evidence_coverage,
               created_at, report_text, assessment_payload
        FROM assessments
        WHERE user_id = ?
        ORDER BY datetime(created_at) DESC
        """,
        conn,
        params=(user_id,)
    )
    conn.close()
    return df


def get_all_users():
    conn = get_conn()
    df = pd.read_sql_query(
        """
        SELECT id, full_name, email, role, created_at
        FROM users
        ORDER BY datetime(created_at) DESC
        """,
        conn
    )
    conn.close()
    return df


def get_all_assessments():
    conn = get_conn()
    df = pd.read_sql_query(
        """
        SELECT id, user_id, institution_name, discipline, programme_name,
               assessment_type, total_points, programme_score, rule_based_status,
               predicted_status, ml_confidence, evidence_coverage, extraction_summary, created_at
        FROM assessments
        ORDER BY datetime(created_at) DESC
        """,
        conn
    )
    conn.close()
    return df

# =========================================================
# MODEL TRAINING
# =========================================================
@st.cache_resource
def train_random_forest():
    df = pd.read_csv("nuc_dataset_22150_questionnaire_full.csv")

    target = "actual_accreditation_status"
    drop_cols = [
        "programme_id",
        "institution_name",
        "programme_name",
        "self_study_score",
        "academic_score",
        "staffing_score",
        "physical_facilities_score",
        "library_score",
        "funding_score",
        "research_score",
        "programme_score",
        target,
    ]

    X = df.drop(columns=drop_cols, errors="ignore")
    y = df[target]

    le = LabelEncoder()
    y = le.fit_transform(y)

    X = pd.get_dummies(X, drop_first=True)
    training_columns = X.columns

    X_train, _, y_train, _ = train_test_split(
        X, y, test_size=0.2, random_state=42, stratify=y
    )

    rf = RandomForestClassifier(
        n_estimators=300,
        random_state=42,
        class_weight="balanced",
        n_jobs=-1,
    )
    rf.fit(X_train, y_train)
    metadata = {
        "model_name": "RandomForestClassifier",
        "model_version": "rf-2026-07-13",
        "training_date": datetime.utcnow().date().isoformat(),
        "classes": list(le.classes_),
    }
    return rf, le, training_columns, metadata


def load_random_forest():
    return train_random_forest()


def prepare_model_input(input_data: dict, training_columns):
    input_df = pd.DataFrame([input_data])
    input_encoded = pd.get_dummies(input_df)
    return input_encoded.reindex(columns=training_columns, fill_value=0)


def validate_model_features(input_encoded, training_columns):
    return list(input_encoded.columns) == list(training_columns)


def predict_accreditation(model, label_encoder, input_encoded):
    prediction = model.predict(input_encoded)
    return label_encoder.inverse_transform(prediction)[0]


def predict_probabilities(model, label_encoder, input_encoded):
    probabilities = model.predict_proba(input_encoded)[0]
    return {
        label: float(probabilities[idx])
        for idx, label in enumerate(label_encoder.classes_)
    }


model, label_encoder, training_columns, model_metadata = load_random_forest()

# =========================================================
# OPENAI CLIENT
# =========================================================
client = None
if "OPENAI_API_KEY" in st.secrets:
    clean_key = st.secrets["OPENAI_API_KEY"]
    clean_key = clean_key.replace("\u200b", "").strip()
    clean_key = re.sub(r"[^\x00-\x7F]+", "", clean_key)
    if clean_key:
        client = OpenAI(api_key=clean_key)

# =========================================================
# SESSION STATE
# =========================================================
if "user" not in st.session_state:
    st.session_state.user = None

if "page" not in st.session_state:
    st.session_state.page = "Login"

# =========================================================
# BRANDING AND UI HELPERS
# =========================================================
ASSET_DIR = Path("assets")
IMAGE_DIR = ASSET_DIR / "images"
LOGO_PATH = ASSET_DIR / "logo" / "app_logo.svg"
DISCLAIMER_TEXT = (
    "This is an independent accreditation-readiness decision-support platform "
    "and does not issue official NUC accreditation decisions."
)


def safe_image(path, caption=None):
    image_path = Path(path)
    if image_path.exists():
        st.image(str(image_path), caption=caption, width="stretch")


def render_logo(lockup=True):
    if LOGO_PATH.exists():
        st.image(str(LOGO_PATH), width=76)
    if lockup:
        st.markdown(
            """
            <div style="font-weight:800;color:#0b1f3a;line-height:1.15;">
                NUC Accreditation<br>Readiness System
            </div>
            """,
            unsafe_allow_html=True,
        )


def user_initials(full_name: str) -> str:
    parts = [part for part in full_name.strip().split() if part]
    if not parts:
        return "US"
    return "".join(part[0].upper() for part in parts[:2])


def page_intro(title: str, subtitle: str, image_name: str = None):
    safe_title = escape(title)
    safe_subtitle = escape(subtitle)
    left, right = st.columns([2.2, 1])
    with left:
        st.markdown(
            f"""
            <div class="hero-card">
                <h1 style="margin:0 0 .35rem;">{safe_title}</h1>
                <p style="margin:0;max-width:760px;">{safe_subtitle}</p>
            </div>
            """,
            unsafe_allow_html=True,
        )
    with right:
        if image_name:
            safe_image(IMAGE_DIR / image_name)


def render_empty_state(image_name: str, title: str, body: str):
    safe_title = escape(title)
    safe_body = escape(body)
    left, right = st.columns([1, 2])
    with left:
        safe_image(IMAGE_DIR / image_name)
    with right:
        st.markdown(
            f"""
            <div class="institution-card">
                <h3 class="section-title">{safe_title}</h3>
                <p class="subtle-copy">{safe_body}</p>
            </div>
            """,
            unsafe_allow_html=True,
        )


def render_assessment_card(title: str, body: str, image_name: str, button_label: str, route: str):
    st.markdown('<div class="institution-card">', unsafe_allow_html=True)
    safe_image(IMAGE_DIR / image_name)
    st.markdown(f"### {escape(title)}")
    st.write(body)
    if st.button(button_label, width="stretch", key=f"route_{route}"):
        st.session_state.assessment_route = route
        st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)

# =========================================================
# MASTER DATA
# =========================================================
institutions = KNOWN_INSTITUTIONS

discipline_programmes = {
    "Computing": [
        "Computer Science", "Software Engineering", "Information Technology",
        "Cybersecurity", "Data Science"
    ],
    "Education": [
        "Educational Management", "Guidance and Counselling",
        "Curriculum Studies", "Adult Education"
    ],
    "Engineering": [
        "Civil Engineering", "Mechanical Engineering", "Electrical Engineering",
        "Chemical Engineering", "Petroleum Engineering", "Mechatronics Engineering",
        "Biomedical Engineering", "Agricultural Engineering"
    ],
    "Management": [
        "Accounting", "Business Administration", "Economics",
        "Banking and Finance", "Public Administration"
    ],
    "Science": [
        "Biology", "Chemistry", "Physics", "Microbiology",
        "Biochemistry", "Mathematics", "Statistics", "Geology"
    ],
}

disciplines = list(discipline_programmes.keys())

# =========================================================
# SECTION MAXIMUMS
# =========================================================
SECTION_MAX = {
    "Academic Content": 37,
    "Staffing": 35,
    "Physical Facilities": 27,
    "Library": 18,
    "Funding": 3,
    "Research and Collaboration": 3,
    "Tracer and Employers’ Rating": 3,
}
TOTAL_MAX_SCORE = 125

# =========================================================
# QUESTION OPTION BANKS
# =========================================================
PHILOSOPHY_OPTIONS = [
    {"label": "Clearly defined and similar to those laid down in the CCMAS for the programme", "score": 1.0},
    {"label": "Not well stated/Not in line with those laid down in the CCMAS for the programme", "score": 0.0},
]

RELATIONSHIP_CCMAS_OPTIONS = [
    {"label": "The curriculum is adequate for the degree programme as it contains all the core/compulsory courses prescribed by the CCMAS", "score": 1.0},
    {"label": "Not fully adequate for the degree programme as a few core/compulsory courses are omitted in some levels of study as prescribed by the CCMAS", "score": 0.5},
    {"label": "Not adequate for the degree programme as it does not contain all the core/compulsory courses across all the levels of study as prescribed by the CCMAS", "score": 0.0},
]

INNOVATION_OPTIONS = [
    {"label": "Inclusion of four or more additional/innovative courses in the curriculum", "score": 1.0},
    {"label": "Inclusion of less than four innovative courses in the curriculum", "score": 0.5},
    {"label": "Non-inclusion of additional/innovative courses in the curriculum", "score": 0.0},
]

COVERAGE_OPTIONS = [
    {"label": "The curriculum is adequately covered", "score": 1.0},
    {"label": "The curriculum is substantially covered", "score": 0.5},
    {"label": "The curriculum is not adequately covered", "score": 0.0},
]

ADMISSION_OPTIONS = [
    {"label": "All students enrolled in the programme meet the admission requirements", "score": 1.0},
    {"label": "Some of the students enrolled in the programme did not meet the admission requirements", "score": 0.0},
]

ACADEMIC_REG_OPTIONS = [
    {"label": "Available, quite clear, are in use and well publicized to students", "score": 1.0},
    {"label": "Available, not clear, but in use and well publicized to students", "score": 0.5},
    {"label": "Not available", "score": 0.0},
]

TEST_EXAM_OPTIONS = [
    {"label": "Very good standard and quality and adequately cover the curriculum", "score": 1.0},
    {"label": "Good standard and quality but do not cover the curriculum", "score": 0.5},
    {"label": "Not of good standard and do not adequately cover the curriculum", "score": 0.0},
]

EVALUATION_WORK_OPTIONS = [
    {"label": "Marking schemes exist, are well-developed and the grading of projects, CA, course work and exam scripts is consistent", "score": 1.0},
    {"label": "Marking schemes exist, not well-developed and the grading of projects, CA, course work and exam scripts is not consistent", "score": 0.5},
    {"label": "Marking schemes do not exist, and the grading of projects, CA, course work and exam script is poor and not consistent", "score": 0.0},
]

DEGREE_PROJECT_OPTIONS = [
    {"label": "Good quality, well supervised, innovative topics and creative", "score": 1.0},
    {"label": "Innovative and creative but not well supervised", "score": 0.5},
    {"label": "Not well supervised, lack creativity and not innovative", "score": 0.0},
]

PRACTICAL_WORK_OPTIONS = [
    {"label": "Practicals conducted, good quality, depth and scope covered", "score": 1.0},
    {"label": "Practicals conducted, good quality but did not cover the curriculum", "score": 0.66},
    {"label": "Practicals conducted but not of good quality", "score": 0.33},
    {"label": "Practical not conducted", "score": 0.0},
]

STUDENTS_COURSE_EVAL_OPTIONS = [
    {"label": "The course content, learning materials, course delivery, physical facilities are adequate", "score": 1.0},
    {"label": "The course content, learning materials, course delivery, physical facilities are fairly adequate", "score": 0.5},
    {"label": "The course content, learning materials, course delivery, physical facilities are not adequate", "score": 0.0},
]

SKILLS_ACQUISITION_OPTIONS = [
    {"label": "Students demonstrate very good level of hard and soft skills relevant to the programme", "score": 1.0},
    {"label": "Students demonstrate good level of hard and soft skills relevant to the programme", "score": 0.66},
    {"label": "Students demonstrate average level of hard and soft skills relevant to the programme", "score": 0.33},
    {"label": "Students demonstrate poor level of hard and soft skills relevant to the programme", "score": 0.0},
]

EXTERNAL_EXAM_OPTIONS = [
    {"label": "External examiners system exists. Qualified assessors are engaged and the work done is of good standard", "score": 1.0},
    {"label": "External examiners system exists. Qualified assessors are engaged but the work done is not of good standard", "score": 0.5},
    {"label": "External examiners system exists but the quality of assessors engaged is poor OR external examiners system does not exist", "score": 0.0},
]

IQA_OPTIONS = [
    {"label": "Exists and effective", "score": 1.0},
    {"label": "Exists not effective", "score": 0.5},
    {"label": "Does not exist", "score": 0.0},
]

COMPETENCE_OPTIONS = [
    {"label": "Competent", "score": 1.0},
    {"label": "Not competent", "score": 0.0},
]

ADMIN_OPTIONS = [
    {"label": "Run by a qualified academic staff (Senior Lecturer and above) and very effective and efficient", "score": 1.0},
    {"label": "Run by a qualified academic staff (Senior Lecturer and above) and efficient", "score": 0.5},
    {"label": "Run by an inexperienced academic and generally ineffective and inefficient", "score": 0.0},
]

LAB_SPACE_OPTIONS = [
    {"label": "Adequate and meets the NUC space standard by 70% or more", "score": 1.0},
    {"label": "Meets 60% but less than 70% of the NUC space standards", "score": 0.66},
    {"label": "Meets 50% but less than 60% of the NUC space standards", "score": 0.33},
    {"label": "Meets less than 50% of the NUC space standards", "score": 0.0},
]

LAB_EQUIPMENT_OPTIONS = [
    {"label": "80% or more", "score": 1.0},
    {"label": "70% but less than 80%", "score": 0.75},
    {"label": "60% but less than 70%", "score": 0.5},
    {"label": "50% but less than 60%", "score": 0.25},
    {"label": "Less than 50%", "score": 0.0},
]

CLASSROOM_SPACE_OPTIONS = [
    {"label": "70% or more", "score": 1.0},
    {"label": "60% but less than 70%", "score": 0.66},
    {"label": "50% but less than 60%", "score": 0.33},
    {"label": "Less than 50%", "score": 0.0},
]

CLASSROOM_EQUIPMENT_OPTIONS = [
    {"label": "Adequate and well maintained", "score": 1.0},
    {"label": "Adequate but not well maintained/slightly inadequate but well maintained", "score": 0.5},
    {"label": "Inadequate and not well maintained", "score": 0.0},
]

OFFICE_OPTIONS = [
    {"label": "Adequate in space and well equipped", "score": 1.0},
    {"label": "Slightly inadequate in space but well equipped", "score": 0.66},
    {"label": "Adequate in space but ill-equipped OR inadequate in space but well equipped", "score": 0.33},
    {"label": "Inadequate in space and ill-equipped or inappropriate", "score": 0.0},
]

SAFETY_OPTIONS = [
    {"label": "Safe and comply with all Government Laws relating to fire and environmental sanitation including adequate and functional toilet facilities", "score": 1.0},
    {"label": "Reasonably safe and comply with most Government Laws relating to fire and environmental sanitation including some functional toilet facilities", "score": 0.5},
    {"label": "Unsafe and violate Government Laws relating to fire and environmental sanitation including toilet facilities", "score": 0.0},
]

LIBRARY_HOLDINGS_OPTIONS = [
    {"label": "Good quality and very relevant and adequate in number and coverage", "score": 1.0},
    {"label": "Good quality and adequate in number and coverage", "score": 0.66},
    {"label": "Inadequate in number, quality and coverage", "score": 0.33},
    {"label": "Poor quality and inadequate in number and coverage", "score": 0.0},
]

CURRENCY_OPTIONS = [
    {"label": "Very current for both recommended text and journals", "score": 1.0},
    {"label": "Very current for recommended text but current for journals or vice versa", "score": 0.75},
    {"label": "Current for recommended text and journals", "score": 0.5},
    {"label": "Current for recommended text but not for journals or vice versa", "score": 0.25},
    {"label": "Not current at all for both recommended text and journals", "score": 0.0},
]

ACCESS_EBOOKS_OPTIONS = [
    {"label": "Internet, Library website, Campus Hotspots and Intranet", "score": 1.0},
    {"label": "Campus Hotspots and Intranet", "score": 0.66},
    {"label": "Intranet", "score": 0.33},
    {"label": "No access media", "score": 0.0},
]

SUBSCRIPTION_OPTIONS = [
    {"label": "Active subscription to at least 2 relevant databases in addition to open source materials", "score": 1.0},
    {"label": "Active subscription to at least 1 relevant database in addition to open source materials", "score": 0.75},
    {"label": "Access to only open source materials", "score": 0.5},
    {"label": "Access to only offline materials", "score": 0.25},
    {"label": "No subscription and no access to offline and open source materials", "score": 0.0},
]

RESEARCH_COLLAB_OPTIONS = [
    {"label": "Available, multidisciplinary approach engaged", "score": 1.0},
    {"label": "Available, no multidisciplinary approach, result applied", "score": 0.66},
    {"label": "Available only", "score": 0.33},
    {"label": "Not available at all", "score": 0.0},
]

TRACER_OPTIONS = [
    {"label": "Tracer system is in place and graduates’ performance on the job is very good", "score": 1.0},
    {"label": "Tracer system not in place but graduates’ performance on the job is good or vice versa", "score": 0.5},
    {"label": "Tracer system not in place and performance below average", "score": 0.0},
]

TRACER_NO_GRAD_OPTIONS = [
    {"label": "Very good", "score": 1.0},
    {"label": "Average", "score": 0.5},
    {"label": "Below average", "score": 0.0},
]

# =========================================================
# HELPERS
# =========================================================
def ask_scored_question(label: str, options: list, key: str):
    option_labels = [opt["label"] for opt in options]
    selected = st.selectbox(
        label,
        option_labels,
        index=None,
        placeholder="Select an option",
        key=key,
    )
    if selected is None:
        return None
    for opt in options:
        if opt["label"] == selected:
            return opt["score"]
    return None


def pct(numerator: int, denominator: int) -> float:
    return 0.0 if denominator <= 0 else (numerator / denominator) * 100


def display_score_card(title: str, points, max_points: float):
    if points is None:
        text = "—"
    else:
        text = f"{points:.2f} / {max_points}"
    st.markdown(
        f'<div class="score-card"><h4>{title}</h4><h2>{text}</h2></div>',
        unsafe_allow_html=True
    )


def get_feature_importance(_model, _training_columns):
    try:
        importances = _model.feature_importances_
    except Exception:
        return pd.DataFrame(columns=["feature", "importance"])

    return pd.DataFrame({
        "feature": _training_columns,
        "importance": importances
    }).sort_values(by="importance", ascending=False)


def inputs_complete(values):
    return all(v is not None for v in values)


def compute_section_points(responses: dict, max_points: float):
    vals = list(responses.values())
    if not vals or any(v is None for v in vals):
        return None
    return (sum(vals) / len(vals)) * max_points


def classify_funding_amount(amount):
    if amount is None:
        return None, None
    if amount >= 50_000_000:
        return "Very adequate", 1.0
    if amount >= 20_000_000:
        return "Adequate", 0.66
    if amount >= 5_000_000:
        return "Inadequate", 0.33
    return "Poor", 0.0


def determine_nuc_status(programme_score):
    if programme_score is None:
        return "UNDETERMINED"
    if programme_score >= 70:
        return "FULL"
    if programme_score >= 60:
        return "INTERIM"
    return "DENIED"


@dataclass
class FieldEvidence:
    field_name: str
    extracted_value: object
    confidence: float
    source_file: str
    source_page_or_section: str
    evidence_text: str
    status: str
    user_confirmed: bool = False
    extraction_method: str = "text_regex"


def read_uploaded_file_inline_legacy(uploaded_file):
    suffix = uploaded_file.name.split(".")[-1].lower()
    data = uploaded_file.getvalue()
    warnings = []
    text_blocks = []
    tables = []

    if suffix == "pdf":
        if PdfReader is None:
            warnings.append("Install pypdf to extract PDF content.")
        else:
            reader = PdfReader(BytesIO(data))
            for index, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_blocks.append({"section": f"Page {index}", "text": page_text})
            if not text_blocks:
                warnings.append("No meaningful extractable text found. OCR was not used.")
    elif suffix == "docx":
        if Document is None:
            warnings.append("Install python-docx to extract DOCX content.")
        else:
            document = Document(BytesIO(data))
            text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
            text_blocks.append({"section": "Document body", "text": text})
            for table in document.tables:
                tables.append([[cell.text for cell in row.cells] for row in table.rows])
    elif suffix == "xlsx":
        try:
            sheets = pd.read_excel(BytesIO(data), sheet_name=None)
            for sheet_name, sheet_df in sheets.items():
                tables.append(sheet_df.fillna("").astype(str).values.tolist())
                text_blocks.append({"section": sheet_name, "text": sheet_df.to_csv(index=False)})
        except Exception as exc:
            warnings.append(f"Could not extract XLSX content: {exc}")
    elif suffix == "txt":
        text_blocks.append({"section": "Text file", "text": data.decode("utf-8", errors="ignore")})
    else:
        warnings.append("Unsupported file type.")

    combined = "\n".join(block["text"] for block in text_blocks).lower()
    detected = "unknown supporting evidence"
    for category, keywords in DOCUMENT_CATEGORIES.items():
        if any(keyword in combined for keyword in keywords):
            detected = category.replace("_", " ")
            break

    return {
        "file_name": uploaded_file.name,
        "file_type": suffix,
        "text_blocks": text_blocks,
        "tables": tables,
        "detected_document_type": detected,
        "extraction_warnings": warnings,
        "ocr_used": False,
    }


def find_number_evidence(parsed_docs, field_name, patterns, confidence=0.82):
    found = []
    for doc in parsed_docs:
        for block in doc["text_blocks"]:
            for pattern in patterns:
                match = re.search(pattern, block["text"], flags=re.IGNORECASE)
                if match:
                    value = int(match.group(1).replace(",", ""))
                    found.append(FieldEvidence(
                        field_name=field_name,
                        extracted_value=value,
                        confidence=confidence,
                        source_file=doc["file_name"],
                        source_page_or_section=block["section"],
                        evidence_text=match.group(0)[:240],
                        status="confirmed",
                    ))
    if not found:
        return FieldEvidence(field_name, None, 0.0, "", "", "", "missing")

    unique_values = {item.extracted_value for item in found}
    if len(unique_values) > 1:
        first = found[0]
        first.status = "conflicting"
        first.confidence = min(first.confidence, 0.55)
        first.evidence_text = "Conflicting values found: " + ", ".join(str(v) for v in sorted(unique_values))
        return first
    return found[0]


def extract_document_evidence(parsed_docs):
    field_patterns = {
        "academic_staff_count": [r"(?:academic staff|teaching staff)[^\d]{0,40}(\d{1,4})"],
        "student_count": [r"(?:student enrolment|number of students|students enrolled)[^\d]{0,40}(\d{1,5})"],
        "core_staff_count": [r"(?:core staff|staff core to the subject)[^\d]{0,40}(\d{1,4})"],
        "professors_readers_count": [r"(?:professors?/readers?|professor and reader)[^\d]{0,40}(\d{1,4})"],
        "senior_lecturers_count": [r"(?:senior lecturers?)[^\d]{0,40}(\d{1,4})"],
        "lecturer_one_and_below_count": [r"(?:lecturer i and below|lecturers? one and below|lecturer 1 and below)[^\d]{0,40}(\d{1,4})"],
        "phd_holders_count": [r"(?:ph\.?d holders?|doctorate holders?)[^\d]{0,40}(\d{1,4})"],
        "non_teaching_staff_count": [r"(?:non-teaching staff|non academic staff|non-academic staff)[^\d]{0,40}(\d{1,4})"],
        "academic_staff_development_count": [r"(?:academic staff development|academic staff trained)[^\d]{0,40}(\d{1,4})"],
        "non_academic_staff_development_count": [r"(?:non-academic staff development|non academic staff trained)[^\d]{0,40}(\d{1,4})"],
    }
    return {
        field: find_number_evidence(parsed_docs, field, patterns)
        for field, patterns in field_patterns.items()
    }


def evidence_to_groups(evidence):
    groups = {"confirmed": [], "needs_review": [], "conflicting": [], "missing": []}
    for item in evidence.values():
        groups.setdefault(item.status, []).append(item)
    return groups


def evidence_coverage(evidence):
    if not evidence:
        return 0.0
    resolved = [item for item in evidence.values() if item.status == "confirmed" and item.extracted_value is not None]
    return round((len(resolved) / len(evidence)) * 100, 2)


def build_model_input_from_scores(base_inputs):
    return prepare_model_input(base_inputs, training_columns)


def keyword_score(parsed_docs, keywords, label):
    text = "\n".join(
        block["text"].lower()
        for doc in parsed_docs
        for block in doc["text_blocks"]
    )
    hits = [kw for kw in keywords if kw.lower() in text]
    if len(hits) >= 2:
        return 1.0, f"{label}: evidence found for {', '.join(hits[:3])}"
    if hits:
        return 0.5, f"{label}: partial evidence found for {hits[0]}"
    return None, f"{label}: no supporting evidence found"


def render_evidence_group(title, items):
    st.markdown(f"#### {title}")
    if not items:
        st.caption("No items in this group.")
        return
    rows = []
    for item in items:
        rows.append({
            "Field": item.field_name,
            "Value": item.extracted_value,
            "Confidence": item.confidence,
            "Source": item.source_file,
            "Location": item.source_page_or_section,
            "Evidence": item.evidence_text,
            "Status": item.status,
        })
    st.dataframe(pd.DataFrame(rows), width="stretch")


def get_document_signature(programme, parsed_docs):
    parts = [programme or ""]
    parts.extend(f"{doc.get('file_name')}:{len(doc.get('text_blocks', []))}:{len(doc.get('tables', []))}" for doc in parsed_docs)
    return "|".join(parts)


def render_institution_identification(parsed_docs):
    extraction = extract_institution_name(parsed_docs, institutions)
    signature = get_document_signature("computing_document", parsed_docs)
    state_key = f"institution_resolution_{signature}"
    state = st.session_state.setdefault(state_key, {
        "confirmed_value": None,
        "corrected_by_user": False,
        "correction_timestamp": None,
    })

    display_value = extraction.matched_known_institution or extraction.normalised_candidate
    st.markdown("#### Institution Identified")

    if extraction.status == "confirmed":
        st.success(display_value)
        st.caption(f"Source: {extraction.source_file} — {extraction.source_page_or_section}")
        st.caption(f"Confidence: {extraction.confidence:.0%}")
        c1, c2 = st.columns(2)
        with c1:
            if st.button("Confirm Institution", type="primary", width="stretch", key=f"confirm_{state_key}"):
                state["confirmed_value"] = display_value
                state["corrected_by_user"] = False
                state["correction_timestamp"] = None
        with c2:
            edit_clicked = st.button("Edit Institution", width="stretch", key=f"edit_{state_key}")
        if edit_clicked or state.get("edit_mode"):
            state["edit_mode"] = True
            corrected = st.selectbox(
                "Correct institution",
                institutions + ["Other institution"],
                index=None,
                placeholder="Select the correct institution",
                key=f"correct_select_{state_key}",
            )
            if corrected == "Other institution":
                corrected = st.text_input("Other institution name", key=f"correct_other_{state_key}")
            if corrected and st.button("Save Institution Correction", width="stretch", key=f"save_correct_{state_key}"):
                state["confirmed_value"] = corrected
                state["corrected_by_user"] = corrected != display_value
                state["correction_timestamp"] = datetime.utcnow().isoformat()
                state["edit_mode"] = False

    elif extraction.status == "needs_review":
        st.warning("Review and confirm the institution before continuing.")
        default_options = institutions + ["Other institution"]
        default_index = institutions.index(display_value) if display_value in institutions else None
        selected = st.selectbox(
            "Institution name",
            default_options,
            index=default_index,
            placeholder="Select the correct institution",
            key=f"review_select_{state_key}",
        )
        if selected == "Other institution":
            selected = st.text_input("Other institution name", value=display_value if display_value not in institutions else "", key=f"review_other_{state_key}")
        if selected and st.button("Confirm Institution", type="primary", width="stretch", key=f"review_confirm_{state_key}"):
            state["confirmed_value"] = selected
            state["corrected_by_user"] = selected != display_value
            state["correction_timestamp"] = datetime.utcnow().isoformat() if selected != display_value else None

    elif extraction.status == "conflicting":
        st.error("Conflicting institution names were found. Select or enter the correct institution before continuing.")
        rows = []
        for candidate in extraction.alternative_candidates:
            rows.append({
                "Institution": candidate.get("matched_known_institution") or candidate.get("normalised_candidate"),
                "Source": candidate.get("source_file"),
                "Location": candidate.get("source_page_or_section"),
                "Confidence": candidate.get("confidence"),
                "Evidence": candidate.get("evidence_text"),
            })
        st.dataframe(pd.DataFrame(rows), width="stretch")
        selected = st.selectbox(
            "Correct institution",
            institutions + ["Other institution"],
            index=None,
            placeholder="Select the correct institution",
            key=f"conflict_select_{state_key}",
        )
        if selected == "Other institution":
            selected = st.text_input("Other institution name", key=f"conflict_other_{state_key}")
        if selected and st.button("Confirm Institution", type="primary", width="stretch", key=f"conflict_confirm_{state_key}"):
            state["confirmed_value"] = selected
            state["corrected_by_user"] = selected != display_value
            state["correction_timestamp"] = datetime.utcnow().isoformat()

    else:
        st.warning("We could not reliably identify the institution from the uploaded documents. Please select or enter the institution name.")
        selected = st.selectbox(
            "Institution name",
            institutions + ["Other institution"],
            index=None,
            placeholder="Search or select institution",
            key=f"missing_select_{state_key}",
        )
        if selected == "Other institution":
            selected = st.text_input("Other institution name", key=f"missing_other_{state_key}")
        if selected and st.button("Confirm Institution", type="primary", width="stretch", key=f"missing_confirm_{state_key}"):
            state["confirmed_value"] = selected
            state["corrected_by_user"] = bool(display_value and selected != display_value)
            state["correction_timestamp"] = datetime.utcnow().isoformat()

    confirmed = state.get("confirmed_value")
    if confirmed:
        st.info(f"Final institution for this assessment: {confirmed}")

    evidence_record = extraction.to_field_evidence(user_confirmed=bool(confirmed))
    evidence_record.update({
        "confirmed_value": confirmed,
        "corrected_by_user": bool(state.get("corrected_by_user")),
        "correction_timestamp": state.get("correction_timestamp"),
    })
    return confirmed, extraction, evidence_record


def render_computing_document_assessment():
    st.subheader("Computing Document Assessment")
    st.caption("Upload Computing accreditation documents. Original files are parsed in memory and are not persisted by default.")

    programme = st.selectbox("Computing Programme", COMPUTING_PROGRAMMES, index=None, placeholder="Select Computing programme")
    uploaded_files = st.file_uploader(
        "Upload supporting documents",
        type=["pdf", "docx", "txt", "xlsx"],
        accept_multiple_files=True,
    )

    if not uploaded_files:
        render_empty_state(
            "empty_documents.svg",
            "No documents uploaded",
            "Upload one or more PDF, DOCX, TXT, or XLSX evidence files to begin extraction and review."
        )
        return

    parsed_docs = []
    with st.spinner("Parsing uploaded documents..."):
        for uploaded_file in uploaded_files:
            parsed_docs.append(read_uploaded_file(uploaded_file))

    st.markdown("#### Uploaded Files")
    st.dataframe(pd.DataFrame([
        {
            "File": doc["file_name"],
            "Type": doc["file_type"],
            "Detected Document Type": doc["detected_document_type"],
            "Text Blocks": len(doc["text_blocks"]),
            "Tables": len(doc["tables"]),
            "Warnings": "; ".join(doc["extraction_warnings"]),
        }
        for doc in parsed_docs
    ]), width="stretch")

    institution, institution_extraction, institution_evidence = render_institution_identification(parsed_docs)

    evidence = extract_document_evidence(parsed_docs)
    evidence["institution_name"] = FieldEvidence(
        field_name="institution_name",
        extracted_value=institution_evidence.get("extracted_value"),
        confidence=float(institution_evidence.get("confidence") or 0.0),
        source_file=institution_evidence.get("source_file") or "",
        source_page_or_section=institution_evidence.get("source_page_or_section") or "",
        evidence_text=institution_evidence.get("evidence_text") or "",
        status="confirmed" if institution else institution_evidence.get("status", "missing"),
        user_confirmed=bool(institution),
        extraction_method=institution_evidence.get("extraction_method") or "institution_extractor",
    )
    groups = evidence_to_groups(evidence)
    coverage = evidence_coverage(evidence)

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Evidence Coverage", f"{coverage:.2f}%")
    c2.metric("Confirmed", len(groups["confirmed"]))
    c3.metric("Needs Review / Conflicts", len(groups["needs_review"]) + len(groups["conflicting"]))
    c4.metric("Missing", len(groups["missing"]))

    with st.expander("Review Extracted Information", expanded=True):
        render_evidence_group("Confirmed Information", groups["confirmed"])
        render_evidence_group("Information Needing Review", groups["needs_review"])
        render_evidence_group("Conflicting Information", groups["conflicting"])
        render_evidence_group("Missing Required Information", groups["missing"])

    st.markdown("#### Resolve Missing or Uncertain Fields")
    resolved_values = {}
    for field, item in evidence.items():
        if field == "institution_name":
            continue
        if item.status == "confirmed" and item.extracted_value is not None:
            resolved_values[field] = int(item.extracted_value)
            continue
        resolved_values[field] = st.number_input(
            field.replace("_", " ").title(),
            min_value=0,
            value=None,
            placeholder="Enter verified value from institutional records",
            key=f"doc_resolve_{field}",
        )

    has_graduated_students = st.selectbox(
        "Has the programme graduated students?",
        ["Yes", "No"],
        index=None,
        placeholder="Select an option",
        key="doc_graduated_students",
    )

    mandatory_complete = (
        all(value is not None for value in resolved_values.values())
        and programme
        and institution_is_resolved(institution)
        and has_graduated_students
    )
    if not mandatory_complete:
        st.warning("Resolve all mandatory staffing/enrolment fields, programme identity and institution identity before prediction.")
        return

    ratio_score_raw, actual_ratio, staff_ratio_feature, ratio_max = score_staff_student_ratio(
        resolved_values["academic_staff_count"], resolved_values["student_count"]
    )
    core_score_raw, core_pct, core_staff_feature, core_max = score_core_staff(
        resolved_values["core_staff_count"], resolved_values["academic_staff_count"]
    )
    mix_score_raw, mix_pct_dict, staff_mix_feature, mix_max = score_staff_mix_by_rank(
        resolved_values["professors_readers_count"],
        resolved_values["senior_lecturers_count"],
        resolved_values["lecturer_one_and_below_count"],
    )
    phd_score_raw, phd_pct, phd_feature, phd_max = score_phd_qualification(
        resolved_values["phd_holders_count"], resolved_values["academic_staff_count"]
    )
    acad_dev_score_raw, acad_dev_pct, acad_dev_feature, acad_dev_max = score_academic_staff_dev(
        resolved_values["academic_staff_development_count"], resolved_values["academic_staff_count"]
    )
    non_teach_score_raw, non_teach_feature, non_teach_max = score_non_teaching_staff(
        resolved_values["non_teaching_staff_count"]
    )
    non_acad_dev_score_raw, non_acad_dev_pct, non_acad_dev_feature, non_acad_dev_max = score_non_academic_staff_dev(
        resolved_values["non_academic_staff_development_count"], resolved_values["non_teaching_staff_count"]
    )

    academic_score, academic_note = keyword_score(parsed_docs, ["curriculum", "ccmas", "academic regulations", "examination"], "Academic content")
    facility_score, facility_note = keyword_score(parsed_docs, ["laboratory", "equipment", "internet", "power backup"], "Physical facilities")
    library_score, library_note = keyword_score(parsed_docs, ["library", "e-library", "journal", "database"], "Library")
    funding_score_value, funding_note = keyword_score(parsed_docs, ["funding", "budget", "release"], "Funding")
    research_score_value, research_note = keyword_score(parsed_docs, ["research", "publication", "collaboration"], "Research and collaboration")
    tracer_score_value, tracer_note = keyword_score(parsed_docs, ["tracer", "employer", "graduate"], "Tracer and employer rating")

    inferred_scores = {
        "Academic Content": academic_score,
        "Physical Facilities": facility_score,
        "Library": library_score,
        "Funding": funding_score_value,
        "Research and Collaboration": research_score_value,
        "Tracer and Employers’ Rating": tracer_score_value,
    }
    unresolved_sections = [name for name, score in inferred_scores.items() if score is None]
    section_overrides = {}
    if unresolved_sections:
        st.markdown("#### Resolve Unsupported Section Evidence")
        for section_name in unresolved_sections:
            section_overrides[section_name] = st.selectbox(
                section_name,
                ["Adequate evidence", "Partial evidence", "No adequate evidence"],
                index=None,
                placeholder="Select after reviewing institutional records",
                key=f"doc_section_{section_name}",
            )
        if any(value is None for value in section_overrides.values()):
            st.warning("Resolve unsupported section evidence before prediction. Absence of evidence is not treated as compliance.")
            return

    for section_name, choice in section_overrides.items():
        inferred_scores[section_name] = {
            "Adequate evidence": 1.0,
            "Partial evidence": 0.5,
            "No adequate evidence": 0.0,
        }[choice]

    total_staffing_raw = (
        ratio_score_raw + core_score_raw + mix_score_raw + phd_score_raw +
        non_teach_score_raw + acad_dev_score_raw + non_acad_dev_score_raw
    )
    total_staffing_max = ratio_max + core_max + mix_max + phd_max + non_teach_max + acad_dev_max + non_acad_dev_max
    staffing_points = (total_staffing_raw / total_staffing_max) * SECTION_MAX["Staffing"]
    section_scores = {
        "Academic Content": inferred_scores["Academic Content"] * SECTION_MAX["Academic Content"],
        "Staffing": staffing_points,
        "Physical Facilities": inferred_scores["Physical Facilities"] * SECTION_MAX["Physical Facilities"],
        "Library": inferred_scores["Library"] * SECTION_MAX["Library"],
        "Funding": inferred_scores["Funding"] * SECTION_MAX["Funding"],
        "Research and Collaboration": inferred_scores["Research and Collaboration"] * SECTION_MAX["Research and Collaboration"],
        "Tracer and Employers’ Rating": inferred_scores["Tracer and Employers’ Rating"] * SECTION_MAX["Tracer and Employers’ Rating"],
    }
    total_points = sum(section_scores.values())
    programme_score = (total_points / TOTAL_MAX_SCORE) * 100
    rule_based_status = determine_nuc_status(programme_score)

    input_data = {
        "curriculum_aligned_with_BMAS": inferred_scores["Academic Content"],
        "innovative_courses_present": inferred_scores["Academic Content"],
        "curriculum_coverage_complete": inferred_scores["Academic Content"],
        "admission_requirements_compliant": inferred_scores["Academic Content"],
        "academic_regulations_defined": inferred_scores["Academic Content"],
        "tests_and_examinations_standardized": inferred_scores["Academic Content"],
        "evaluation_methods_clear": inferred_scores["Academic Content"],
        "degree_projects_adequate": inferred_scores["Academic Content"],
        "practical_work_adequate": inferred_scores["Academic Content"],
        "student_course_evaluation_present": inferred_scores["Academic Content"],
        "skills_acquisition_programme": inferred_scores["Academic Content"],
        "external_examiner_system": inferred_scores["Academic Content"],
        "internal_quality_assurance": inferred_scores["Academic Content"],
        "staff_student_ratio_compliant": staff_ratio_feature,
        "proportion_core_staff_sufficient": core_staff_feature,
        "staff_rank_mix_balanced": staff_mix_feature,
        "academic_staff_qualification_high": phd_feature,
        "staff_competence_verified": 1.0 if resolved_values["academic_staff_count"] >= 6 else 0.0,
        "administrative_support_available": non_teach_feature,
        "non_teaching_staff_adequate": non_teach_feature,
        "academic_staff_development_programme": acad_dev_feature,
        "non_academic_staff_development_programme": non_acad_dev_feature,
        "laboratory_space_adequate": inferred_scores["Physical Facilities"],
        "laboratory_equipment_adequate": inferred_scores["Physical Facilities"],
        "classroom_space_adequate": inferred_scores["Physical Facilities"],
        "classroom_equipment_adequate": inferred_scores["Physical Facilities"],
        "office_accommodation_adequate": inferred_scores["Physical Facilities"],
        "safety_environment_present": inferred_scores["Physical Facilities"],
        "library_holdings_adequate": inferred_scores["Library"],
        "library_material_current": inferred_scores["Library"],
        "e_library_subscription_available": inferred_scores["Library"],
        "e_library_access_good": inferred_scores["Library"],
        "programme_funding_adequate": inferred_scores["Funding"],
        "research_collaboration_active": inferred_scores["Research and Collaboration"],
        "research_output_present": inferred_scores["Research and Collaboration"],
        "employer_rating_positive": inferred_scores["Tracer and Employers’ Rating"],
        "tracer_study_available": inferred_scores["Tracer and Employers’ Rating"],
        "discipline": "Computing",
    }

    if st.button("Predict and Generate Computing Report", type="primary", width="stretch"):
        input_encoded = prepare_model_input(input_data, training_columns)
        model_status = predict_accreditation(model, label_encoder, input_encoded)
        probabilities = predict_probabilities(model, label_encoder, input_encoded)
        ml_confidence = probabilities.get(model_status, 0.0)
        computing_findings = evaluate_computing_rules(
            resolved_values,
            {"actual_ratio": actual_ratio, "phd_pct": phd_pct, "mix_pct_dict": mix_pct_dict},
        )
        notes = [academic_note, facility_note, library_note, funding_note, research_note, tracer_note]
        missing_fields = [item.field_name for item in groups["missing"]]
        conflicting_fields = [item.field_name for item in groups["conflicting"]]
        weaknesses = [
            name for name, score in section_scores.items()
            if score < (SECTION_MAX[name] * 0.6)
        ] + computing_findings

        report = build_professional_report({
            "institution": institution,
            "discipline": "Computing",
            "programme": programme,
            "assessment_type": "computing_document",
            "programme_score": programme_score,
            "total_points": total_points,
            "rule_based_status": rule_based_status,
            "ml_predicted_status": model_status,
            "ml_confidence": ml_confidence,
            "probabilities": probabilities,
            "evidence_coverage": coverage,
            "section_scores": section_scores,
            "weaknesses": weaknesses + notes,
            "missing_fields": missing_fields,
            "conflicting_fields": conflicting_fields,
            "computing_findings": computing_findings,
        })

        st.subheader("Section Scores")
        st.dataframe(pd.DataFrame([
            {"Section": name, "Points": round(score, 2), "Maximum": SECTION_MAX[name]}
            for name, score in section_scores.items()
        ]), width="stretch")
        st.subheader("Score-Based Readiness Classification")
        st.success(rule_based_status)
        st.subheader("Random Forest Prediction")
        st.info(f"{model_status} ({ml_confidence:.1%} confidence)")
        if model_status != rule_based_status:
            st.warning("The machine-learning prediction and score-based readiness classification differ. Review the highlighted criteria, supporting evidence and unresolved information before relying on the result.")
        else:
            st.success("The Random Forest prediction agrees with the score-based classification.")

        st.subheader("Accreditation-Readiness Report")
        st.markdown(report)

        payload = {
            "assessment_type": "computing_document",
            "resolved_values": resolved_values,
            "section_scores": section_scores,
            "field_evidence": {field: asdict(item) for field, item in evidence.items()},
            "institution_evidence": institution_evidence,
            "parsed_documents": parsed_docs,
            "model_probabilities": probabilities,
            "computing_findings": computing_findings,
        }
        save_assessment(
            user_id=st.session_state.user["id"],
            institution=institution,
            discipline="Computing",
            programme=programme,
            programme_score=float(programme_score),
            predicted_status=model_status,
            actual_ratio=float(actual_ratio),
            core_pct=float(core_pct),
            phd_pct=float(phd_pct),
            assessment_payload=json.dumps(payload),
            report_text=report,
            assessment_type="computing_document",
            rule_based_status=rule_based_status,
            ml_confidence=float(ml_confidence),
            extraction_summary=json.dumps({
                "files": [doc["file_name"] for doc in parsed_docs],
                "coverage": coverage,
                "institution": institution_evidence,
                "warnings": [warning for doc in parsed_docs for warning in doc["extraction_warnings"]],
            }),
            institution_evidence=json.dumps(institution_evidence),
            extracted_institution_value=institution_evidence.get("extracted_value"),
            confirmed_institution_value=institution,
            institution_corrected_by_user=bool(institution_evidence.get("corrected_by_user")),
            institution_correction_timestamp=institution_evidence.get("correction_timestamp"),
            evidence_coverage=float(coverage),
            total_points=float(total_points),
        )
        st.success("Computing document assessment and report saved successfully.")
        st.download_button("Download Markdown Report", report, file_name="computing_readiness_report.md")
        st.download_button("Download JSON Assessment Record", json.dumps(payload, indent=2), file_name="computing_assessment_record.json")
        st.download_button(
            "Download CSV Score Summary",
            pd.DataFrame([{"section": k, "points": v, "maximum": SECTION_MAX[k]} for k, v in section_scores.items()]).to_csv(index=False),
            file_name="computing_score_summary.csv",
        )

# =========================================================
# STAFFING CALCULATION ENGINE
# =========================================================
def score_staff_student_ratio(staff_count: int, student_count: int):
    if staff_count is None or student_count is None or staff_count <= 0:
        return None, None, None, None
    actual_ratio = round(student_count / staff_count)
    if actual_ratio <= 20:
        raw_score = 4
    elif actual_ratio <= 25:
        raw_score = 3
    elif actual_ratio <= 30:
        raw_score = 2
    elif actual_ratio <= 35:
        raw_score = 1
    else:
        raw_score = 0
    normalized = raw_score / 4
    return raw_score, actual_ratio, normalized, 4


def score_core_staff(core_staff_count: int, staff_count: int):
    if core_staff_count is None or staff_count is None or staff_count <= 0:
        return None, None, None, None
    core_pct = pct(core_staff_count, staff_count)
    if core_pct >= 75:
        raw_score = 6
    elif core_pct >= 60:
        raw_score = 4
    elif core_pct >= 50:
        raw_score = 2
    else:
        raw_score = 0
    normalized = raw_score / 6
    return raw_score, core_pct, normalized, 6


def score_staff_mix_by_rank(prof_count: int, senior_count: int, lect1_below_count: int):
    if prof_count is None or senior_count is None or lect1_below_count is None:
        return None, {"prof_pct": None, "senior_pct": None, "others_pct": None}, None, None

    total = prof_count + senior_count + lect1_below_count
    if total <= 0:
        return None, {"prof_pct": None, "senior_pct": None, "others_pct": None}, None, None

    prof_pct = round(pct(prof_count, total))
    senior_pct = round(pct(senior_count, total))
    others_pct = 100 - prof_pct - senior_pct

    prof_ok = abs(prof_pct - 20) <= 2
    senior_ok = abs(senior_pct - 35) <= 2
    others_ok = abs(others_pct - 45) <= 2
    categories_met = sum([prof_ok, senior_ok, others_ok])

    if categories_met == 3:
        raw_score = 5
    elif categories_met >= 1:
        raw_score = 3
    else:
        raw_score = 0

    normalized = raw_score / 5
    return raw_score, {
        "prof_pct": prof_pct,
        "senior_pct": senior_pct,
        "others_pct": others_pct,
    }, normalized, 5


def score_phd_qualification(phd_count: int, total_academic_staff: int):
    if phd_count is None or total_academic_staff is None or total_academic_staff <= 0:
        return None, None, None, None
    phd_pct = pct(phd_count, total_academic_staff)
    if phd_pct >= 70:
        raw_score = 6
    elif phd_pct >= 60:
        raw_score = 4
    elif phd_pct >= 50:
        raw_score = 2
    else:
        raw_score = 0
    normalized = raw_score / 6
    return raw_score, phd_pct, normalized, 6


def score_academic_staff_dev(trained_count: int, academic_staff_count: int):
    if trained_count is None or academic_staff_count is None or academic_staff_count <= 0:
        return None, None, None, None
    dev_pct = pct(trained_count, academic_staff_count)
    if dev_pct >= 70:
        raw_score = 5
    elif dev_pct >= 60:
        raw_score = 3
    elif dev_pct >= 50:
        raw_score = 1
    else:
        raw_score = 0
    normalized = raw_score / 5
    return raw_score, dev_pct, normalized, 5


def score_non_teaching_staff(non_academic_staff_count: int):
    if non_academic_staff_count is None:
        return None, None, None
    if non_academic_staff_count >= 5:
        return 3, 1.0, 3
    if non_academic_staff_count >= 3:
        return 2, 2/3, 3
    return 0, 0.0, 3


def score_non_academic_staff_dev(trained_count: int, non_academic_staff_count: int):
    if trained_count is None or non_academic_staff_count is None or non_academic_staff_count <= 0:
        return None, None, None, None
    dev_pct = pct(trained_count, non_academic_staff_count)
    if dev_pct >= 70:
        raw_score = 2
    elif dev_pct >= 50:
        raw_score = 1
    else:
        raw_score = 0
    normalized = raw_score / 2
    return raw_score, dev_pct, normalized, 2

# =========================================================
# REPORT CLEANER
# =========================================================
def clean_report_text(text: str) -> str:
    banned_phrases = [
        "Please let me know if you require assistance with specific sections or further elaboration.",
        "Please let me know if you need assistance with specific sections or further elaboration.",
        "Please let me know if you need any further assistance.",
        "Let me know if you need anything else.",
        "Please let me know if you need anything else.",
    ]
    cleaned = text
    for phrase in banned_phrases:
        cleaned = cleaned.replace(phrase, "")
    return cleaned.strip()

# =========================================================
# REPORT GENERATOR
# =========================================================
def build_nuc_style_report(
    institution: str,
    discipline: str,
    programme: str,
    programme_score: float,
    predicted_status: str,
    deficiencies: list,
    remedies: list,
):
    deficiency_lines = "\n".join([f"{idx+1}. {item}" for idx, item in enumerate(deficiencies)]) if deficiencies else "Nil major deficiency identified."
    remedy_lines = "\n".join([f"{idx+1}. {item}" for idx, item in enumerate(remedies)]) if remedies else "Maintain current standards."

    return f"""
## SUMMARY OF PANEL'S REPORT

**DATE OF VISITATION:** {datetime.now().strftime("%d %B %Y")}  
**UNIVERSITY:** {institution}  
**DISCIPLINE:** {discipline}  
**PROGRAMME:** {programme}  

**PROGRAMME SCORE:** {programme_score:.2f}%  
**ACCREDITATION STATUS:** {predicted_status}  

### DEFICIENCIES
{deficiency_lines}

### REMEDIES
{remedy_lines}

## APPROVAL BY THE COMMISSION
Based on further analysis of the report, the Commission upholds the recommendation of **{predicted_status}** accreditation status for the programme.
""".strip()


def evaluate_computing_rules(values, indicators):
    findings = []
    staff_count = values.get("academic_staff_count")
    if staff_count is not None and staff_count < COMPUTING_RULES["minimum_academic_staff"]:
        findings.append("Academic staff count is below the Computing minimum of six.")
    actual_ratio = indicators.get("actual_ratio")
    if actual_ratio is not None and actual_ratio > COMPUTING_RULES["staff_student_ratio_target"]:
        findings.append(f"Staff-student ratio is 1:{actual_ratio}, above the 1:20 target.")
    phd_pct = indicators.get("phd_pct")
    if phd_pct is not None and phd_pct < COMPUTING_RULES["minimum_phd_percentage"]:
        findings.append(f"PhD-qualified staff percentage is {phd_pct:.2f}%, below 70%.")
    mix = indicators.get("mix_pct_dict") or {}
    if all(mix.get(k) is not None for k in ["prof_pct", "senior_pct", "others_pct"]):
        target = COMPUTING_RULES["rank_mix_target"]
        if (
            abs(mix["prof_pct"] - target["professors_readers"]) > target["tolerance"]
            or abs(mix["senior_pct"] - target["senior_lecturers"]) > target["tolerance"]
            or abs(mix["others_pct"] - target["lecturer_one_and_below"]) > target["tolerance"]
        ):
            findings.append("Academic staff rank mix does not meet the 20:35:45 target within ±2 tolerance.")
    return findings


def build_professional_report(context):
    section_scores = context.get("section_scores", {})
    section_lines = "\n".join(
        f"- {name}: {score:.2f} / {SECTION_MAX.get(name, 0)}"
        for name, score in section_scores.items()
        if score is not None
    )
    weaknesses = context.get("weaknesses") or ["No major weakness was detected from resolved evidence."]
    missing = context.get("missing_fields") or []
    conflicts = context.get("conflicting_fields") or []
    computing_findings = context.get("computing_findings") or []
    agreement = (
        "The score-based classification and Random Forest prediction agree."
        if context["rule_based_status"] == context["ml_predicted_status"]
        else "The machine-learning prediction and score-based readiness classification differ. Review the highlighted criteria, supporting evidence and unresolved information before relying on the result."
    )

    return f"""
## Assessment Identity
- Institution: {context['institution']}
- Discipline: {context['discipline']}
- Programme: {context['programme']}
- Assessment Type: {context['assessment_type']}

## Executive Summary
The programme has a readiness score of {context['programme_score']:.2f}% and a score-based classification of {context['rule_based_status']}. The Random Forest prediction is {context['ml_predicted_status']} with {context.get('ml_confidence', 0):.1%} confidence.

## Assessment Method
This advisory assessment combines NUC-style section scoring with a Random Forest model trained on structured accreditation data. It does not grant official NUC accreditation.

## Evidence Coverage Summary
Evidence coverage is {context.get('evidence_coverage', 100):.2f}%. Missing or unresolved evidence is listed below.

## Section Scores
{section_lines}

## Total Points and Programme Score
- Total Points: {context.get('total_points', 0):.2f} / {TOTAL_MAX_SCORE}
- Programme Score: {context['programme_score']:.2f}%

## Score-Based Classification
{context['rule_based_status']}

## Random Forest Prediction
{context['ml_predicted_status']}

## Prediction Confidence
- FULL: {context.get('probabilities', {}).get('FULL', 0):.1%}
- INTERIM: {context.get('probabilities', {}).get('INTERIM', 0):.1%}
- DENIED: {context.get('probabilities', {}).get('DENIED', 0):.1%}

## Agreement or Disagreement Analysis
{agreement}

## High-Impact Weaknesses
{chr(10).join(f'- {item}' for item in weaknesses)}

## Missing Evidence
{chr(10).join(f'- {item}' for item in missing) if missing else '- No mandatory field remains missing.'}

## Conflicting Evidence
{chr(10).join(f'- {item}' for item in conflicts) if conflicts else '- No conflicting mandatory field remains unresolved.'}

## Computing CCMAS Compliance Findings
{chr(10).join(f'- {item}' for item in computing_findings) if computing_findings else '- No Computing-specific violation was detected from resolved fields.'}

## Prioritised Recommendations
{chr(10).join(f'- Resolve: {item}' for item in (weaknesses + missing + conflicts + computing_findings)[:8])}

## Documentation Checklist
- Self-study form
- Curriculum and programme handbook
- Staff list and CV summary
- Course allocation and timetable
- Laboratory and equipment inventory
- Library and e-library evidence
- Research, tracer-study and employer-feedback evidence

## Accreditation-Readiness Action Plan
- Resolve missing and conflicting evidence first.
- Address criteria with low section scores before model-driven presentation improvements.
- Re-run the assessment after evidence and staffing updates are documented.

## Important Advisory Disclaimer
This system provides advisory accreditation-readiness analysis only. Official NUC accreditation decisions can only be made by the appropriate regulatory process.
""".strip()


def generate_advisory_report(
    institution: str,
    programme: str,
    discipline: str,
    programme_score: float,
    predicted_status: str,
    actual_ratio: int,
    ratio_score_raw: int,
    core_pct: float,
    core_score_raw: int,
    mix_score_raw: int,
    phd_pct: float,
    phd_score_raw: int,
    acad_dev_pct: float,
    acad_dev_score_raw: int,
    non_teach_score_raw: int,
    non_acad_dev_pct: float,
    non_acad_dev_score_raw: int,
    top_weaknesses: str,
):
    if client is None:
        deficiencies = [w.strip("- ").strip() for w in top_weaknesses.split("\n") if w.strip()]
        remedies = [f"Improve {item.lower()} before the next accreditation exercise." for item in deficiencies[:5]]
        return build_nuc_style_report(
            institution=institution,
            discipline=discipline,
            programme=programme,
            programme_score=programme_score,
            predicted_status=predicted_status,
            deficiencies=deficiencies[:5],
            remedies=remedies
        )

    prompt = f"""
You are an accreditation advisory expert.

Write a report that mirrors the structure of an NUC accreditation summary report.

Institution: {institution}
Discipline: {discipline}
Programme: {programme}
Programme Score: {programme_score:.2f}%
Predicted Accreditation Status: {predicted_status}

Calculated staffing indicators:
- Staff to Student Ratio: 1:{actual_ratio}
- Staff to Student Ratio Score: {ratio_score_raw}/4
- Core Staff Percentage: {core_pct:.2f}%
- Core Staff Score: {core_score_raw}/6
- Staff Mix by Rank Score: {mix_score_raw}/5
- PhD Holders Percentage: {phd_pct:.2f}%
- Qualifications of Teaching Staff Score: {phd_score_raw}/6
- Academic Staff Development Percentage: {acad_dev_pct:.2f}%
- Academic Staff Development Score: {acad_dev_score_raw}/5
- Non-Teaching Staff Score: {non_teach_score_raw}/3
- Non-Academic Staff Development Percentage: {non_acad_dev_pct:.2f}%
- Non-Academic Staff Development Score: {non_acad_dev_score_raw}/2

High-impact weaknesses:
{top_weaknesses}

Format the output exactly with these headings:
## SUMMARY OF PANEL'S REPORT
## DEFICIENCIES
## REMEDIES
## APPROVAL BY THE COMMISSION

Requirements:
- Make it read like an NUC report
- Be brief, formal, and direct
- List specific deficiencies
- List matching remedies
- State the accreditation status clearly
- Do not add any closing offer for further help
"""

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": "You are an accreditation advisory expert and must return a formal NUC-style accreditation summary report in markdown."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.4,
    )
    return clean_report_text(response.choices[0].message.content)

# =========================================================
# AUTH PAGES
# =========================================================
def render_login():
    st.markdown("### Sign in")
    st.caption("Access assessments, evidence reports, saved records and administrative tools.")
    email = st.text_input("Email", key="login_email")
    password = st.text_input("Password", type="password", key="login_password")

    if st.button("Login", type="primary", width="stretch"):
        user = authenticate_user(email, password)
        if user:
            st.session_state.user = {
                "id": user[0],
                "full_name": user[1],
                "email": user[2],
                "role": user[3],
            }
            st.session_state.page = "Dashboard"
            st.rerun()
        else:
            st.error("Invalid email or password.")

    st.caption("New to the platform?")
    if st.button("Create an account", width="stretch"):
        st.session_state.page = "Signup"
        st.rerun()


def render_signup():
    st.markdown("### Create account")
    st.caption("Set up a user profile for institutional readiness assessments.")
    full_name = st.text_input("Full Name", key="signup_name")
    email = st.text_input("Email", key="signup_email")
    password = st.text_input("Password", type="password", key="signup_password")
    confirm_password = st.text_input("Confirm Password", type="password", key="signup_confirm")

    if st.button("Create Account", type="primary", width="stretch"):
        if not full_name or not email or not password:
            st.error("Please complete all fields.")
        elif password != confirm_password:
            st.error("Passwords do not match.")
        else:
            ok, msg = create_user(full_name, email, password)
            if ok:
                st.success(msg)
                st.session_state.page = "Login"
                st.rerun()
            else:
                st.error(msg)

    st.caption("Already have an account?")
    if st.button("Back to Login", width="stretch"):
        st.session_state.page = "Login"
        st.rerun()

# =========================================================
# PRE-LOGIN
# =========================================================
if st.session_state.user is None:
    if st.session_state.page not in ["Login", "Signup"]:
        st.session_state.page = "Login"

    left, right = st.columns([1.25, 1])
    with left:
        render_logo(lockup=False)
        st.markdown(
            """
            <h1 style="color:#0b1f3a;margin:0.75rem 0 0.35rem;">
                NUC Accreditation Readiness System
            </h1>
            <p class="subtle-copy" style="font-size:1.05rem;max-width:680px;">
                Assess programme readiness, analyse accreditation evidence and generate actionable reports.
            </p>
            <div class="benefit-item"><span class="benefit-dot"></span><span>Structured NUC-style assessment</span></div>
            <div class="benefit-item"><span class="benefit-dot"></span><span>Random Forest readiness prediction</span></div>
            <div class="benefit-item"><span class="benefit-dot"></span><span>Evidence-based advisory reporting</span></div>
            """,
            unsafe_allow_html=True,
        )
        safe_image(IMAGE_DIR / "login_banner.svg")
        st.caption(DISCLAIMER_TEXT)

    with right:
        st.markdown('<div class="auth-panel">', unsafe_allow_html=True)
        render_logo()
        if st.session_state.page == "Signup":
            render_signup()
        else:
            render_login()
        st.markdown("</div>", unsafe_allow_html=True)
    st.stop()

# =========================================================
# SIDEBAR NAV
# =========================================================
with st.sidebar:
    nav_options = ["Dashboard", "New Assessment", "Saved Assessments", "Report History"]
    if st.session_state.user.get("role") == "admin":
        nav_options = ["Dashboard", "New Assessment", "Saved Assessments", "Report History", "Admin Dashboard"]

    nav_labels = {
        "Dashboard": "🏠 Dashboard",
        "New Assessment": "📝 New Assessment",
        "Saved Assessments": "📁 Saved Assessments",
        "Report History": "📊 Report History",
        "Admin Dashboard": "🛡️ Admin Dashboard",
    }
    if (
        "nav_choice" not in st.session_state
        or st.session_state.nav_choice not in nav_options
    ):
        st.session_state.nav_choice = st.session_state.page if st.session_state.page in nav_options else nav_options[0]
        st.session_state.nav_synced_page = st.session_state.nav_choice

    if (
        st.session_state.page in nav_options
        and st.session_state.get("nav_synced_page") != st.session_state.page
    ):
        st.session_state.nav_choice = st.session_state.page
        st.session_state.nav_synced_page = st.session_state.page

    st.markdown('<div class="sidebar-brand">', unsafe_allow_html=True)
    render_logo()
    st.markdown("</div>", unsafe_allow_html=True)

    initials = escape(user_initials(st.session_state.user["full_name"]))
    display_name = escape(st.session_state.user["full_name"])
    display_email = escape(st.session_state.user["email"])
    display_role = escape(st.session_state.user.get("role", "user"))
    st.markdown(
        f"""
        <div class="user-chip">
            <span class="avatar">{initials}</span>
            <div style="display:inline-block;vertical-align:middle;max-width:170px;">
                <div style="font-weight:800;color:#0b1f3a;">{display_name}</div>
                <div style="font-size:.78rem;color:#64748b;overflow-wrap:anywhere;">{display_email}</div>
                <span class="role-badge">{display_role}</span>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    choice = st.radio(
        "Navigation",
        nav_options,
        key="nav_choice",
        format_func=lambda value: nav_labels.get(value, value),
    )
    if choice != st.session_state.page:
        st.session_state.page = choice
        st.session_state.nav_synced_page = choice
    else:
        st.session_state.nav_synced_page = st.session_state.page

    st.divider()
    st.caption(DISCLAIMER_TEXT)
    st.write("")
    if st.button("Logout", width="stretch"):
        st.session_state.user = None
        st.session_state.page = "Login"
        st.rerun()

# =========================================================
# DASHBOARD
# =========================================================
if st.session_state.page == "Dashboard":
    df_assessments = get_user_assessments(st.session_state.user["id"])
    first_name = st.session_state.user["full_name"].split()[0]
    page_intro(
        f"Welcome, {first_name}",
        "Monitor programme readiness, review recent evidence, and move quickly into the next accreditation task.",
        "dashboard_banner.svg",
    )

    q1, q2, q3 = st.columns(3)
    with q1:
        if st.button("Start Manual Assessment", type="primary", width="stretch"):
            st.session_state.assessment_route = "standard_manual"
            st.session_state.page = "New Assessment"
            st.rerun()
    with q2:
        if st.button("Upload Computing Documents", width="stretch"):
            st.session_state.assessment_route = "computing_document"
            st.session_state.page = "New Assessment"
            st.rerun()
    with q3:
        if st.button("View Reports", width="stretch"):
            st.session_state.page = "Report History"
            st.rerun()

    if df_assessments.empty:
        render_empty_state(
            "empty_assessments.svg",
            "No assessments yet",
            "Start a manual questionnaire or upload Computing evidence documents to build your readiness history."
        )
    else:
        avg_score = df_assessments["programme_score"].dropna().mean()
        latest_status = df_assessments.iloc[0]["rule_based_status"] or df_assessments.iloc[0]["predicted_status"]
        incomplete_count = int(df_assessments["programme_score"].isna().sum())
        d1, d2, d3, d4, d5 = st.columns(5)
        d1.metric("Total Assessments", len(df_assessments))
        d2.metric("Average Programme Score", f"{avg_score:.2f}%" if pd.notna(avg_score) else "0.00%")
        d3.metric("Latest Status", latest_status)
        d4.metric("Document Assessments", int((df_assessments["assessment_type"] == "computing_document").sum()))
        d5.metric("Incomplete", incomplete_count)

        chart_data = df_assessments["rule_based_status"].fillna(df_assessments["predicted_status"]).value_counts()
        if not chart_data.empty:
            st.subheader("FULL / INTERIM / DENIED Distribution")
            st.bar_chart(
                chart_data.rename_axis("Status").reset_index(name="Count"),
                x="Status",
                y="Count",
                width="stretch",
            )

        st.subheader("Recent Assessments")
        st.dataframe(df_assessments.head(5), width="stretch")

# =========================================================
# NEW ASSESSMENT
# =========================================================
elif st.session_state.page == "New Assessment":
    page_intro(
        "New Assessment",
        "Choose a readiness workflow, then complete the required evidence and scoring steps.",
        None,
    )
    if "assessment_route" not in st.session_state:
        st.session_state.assessment_route = None

    if st.session_state.assessment_route is None:
        card1, card2 = st.columns(2)
        with card1:
            render_assessment_card(
                "Standard Manual Assessment",
                "Questionnaire-based NUC readiness scoring for Computing, Education, Engineering, Management and Science programmes.",
                "manual_assessment.svg",
                "Start Manual Assessment",
                "standard_manual",
            )
        with card2:
            render_assessment_card(
                "Computing Document Assessment",
                "Upload Computing programme documents, review extracted evidence, resolve gaps, then predict readiness.",
                "document_assessment.svg",
                "Upload Computing Documents",
                "computing_document",
            )
        st.stop()

    top_left, top_right = st.columns([3, 1])
    with top_left:
        st.caption(f"Selected route: {st.session_state.assessment_route}")
    with top_right:
        if st.button("Change Assessment Type", width="stretch"):
            st.session_state.assessment_route = None
            st.rerun()

    if st.session_state.assessment_route == "computing_document":
        render_computing_document_assessment()
        st.stop()

    c1, c2, c3 = st.columns(3)
    with c1:
        institution = st.selectbox("Name of Institution", institutions, index=None, placeholder="Select institution")
    with c2:
        discipline = st.selectbox("Discipline", disciplines, index=None, placeholder="Select discipline")
    with c3:
        available_programmes = discipline_programmes.get(discipline, [])
        programme = st.selectbox("Programme", available_programmes, index=None, placeholder="Select programme")

    st.subheader("Staffing & Enrollment Data")

    s1, s2, s3 = st.columns(3)
    with s1:
        academic_staff_count = st.number_input("Number of academic staff", min_value=1, value=None, placeholder="Enter value")
    with s2:
        student_count = st.number_input("Number of students", min_value=1, value=None, placeholder="Enter value")
    with s3:
        core_staff_count = st.number_input("Number of academic staff core to the subject area", min_value=0, value=None, placeholder="Enter value")

    s4, s5, s6 = st.columns(3)
    with s4:
        professor_reader_count = st.number_input("Number of Professors/Readers", min_value=0, value=None, placeholder="Enter value")
    with s5:
        senior_lecturer_count = st.number_input("Number of Senior Lecturers", min_value=0, value=None, placeholder="Enter value")
    with s6:
        lecturer1_below_count = st.number_input("Number of Lecturers I and below", min_value=0, value=None, placeholder="Enter value")

    s7, s8, s9 = st.columns(3)
    with s7:
        phd_holder_count = st.number_input("Number of Ph.D Holders", min_value=0, value=None, placeholder="Enter value")
    with s8:
        academic_staff_dev_count = st.number_input("Number of academic staff with staff development programme", min_value=0, value=None, placeholder="Enter value")
    with s9:
        non_academic_staff_count = st.number_input("Number of non-teaching staff", min_value=0, value=None, placeholder="Enter value")

    non_academic_staff_dev_count = st.number_input(
        "Number of non-academic staff with staff development programme",
        min_value=0,
        value=None,
        placeholder="Enter value"
    )

    staffing_inputs_complete = inputs_complete([
        academic_staff_count, student_count, core_staff_count,
        professor_reader_count, senior_lecturer_count, lecturer1_below_count,
        phd_holder_count, academic_staff_dev_count,
        non_academic_staff_count, non_academic_staff_dev_count
    ])

    if staffing_inputs_complete:
        ratio_score_raw, actual_ratio, staff_ratio_feature, ratio_max = score_staff_student_ratio(
            academic_staff_count, student_count
        )
        core_score_raw, core_pct, core_staff_feature, core_max = score_core_staff(
            core_staff_count, academic_staff_count
        )
        mix_score_raw, mix_pct_dict, staff_mix_feature, mix_max = score_staff_mix_by_rank(
            professor_reader_count, senior_lecturer_count, lecturer1_below_count
        )
        phd_score_raw, phd_pct, phd_feature, phd_max = score_phd_qualification(
            phd_holder_count, academic_staff_count
        )
        acad_dev_score_raw, acad_dev_pct, acad_dev_feature, acad_dev_max = score_academic_staff_dev(
            academic_staff_dev_count, academic_staff_count
        )
        non_teach_score_raw, non_teach_feature, non_teach_max = score_non_teaching_staff(non_academic_staff_count)
        non_acad_dev_score_raw, non_acad_dev_pct, non_acad_dev_feature, non_acad_dev_max = score_non_academic_staff_dev(
            non_academic_staff_dev_count, non_academic_staff_count
        )

        with st.expander("Computed Staffing Indicators", expanded=True):
            m1, m2, m3 = st.columns(3)
            with m1:
                st.metric("Staff to Student Ratio", f"1 : {actual_ratio}")
            with m2:
                st.metric("Percentage of Staff Core to the Subject Area", f"{round(core_pct, 2)}%")
            with m3:
                st.metric("Ph.D Holders Percentage", f"{round(phd_pct, 2)}%")

            m4, m5, m6 = st.columns(3)
            with m4:
                st.metric(
                    "Staff Mix by Rank",
                    f"{mix_pct_dict['prof_pct']}:{mix_pct_dict['senior_pct']}:{mix_pct_dict['others_pct']}"
                )
            with m5:
                st.metric("Academic Staff Development Percentage", f"{round(acad_dev_pct, 2)}%")
            with m6:
                st.metric("Non-Academic Staff Development Percentage", f"{round(non_acad_dev_pct, 2)}%")
    else:
        ratio_score_raw = actual_ratio = staff_ratio_feature = ratio_max = None
        core_score_raw = core_pct = core_staff_feature = core_max = None
        mix_score_raw = staff_mix_feature = mix_max = None
        mix_pct_dict = {"prof_pct": None, "senior_pct": None, "others_pct": None}
        phd_score_raw = phd_pct = phd_feature = phd_max = None
        acad_dev_score_raw = acad_dev_pct = acad_dev_feature = acad_dev_max = None
        non_teach_score_raw = non_teach_feature = non_teach_max = None
        non_acad_dev_score_raw = non_acad_dev_pct = non_acad_dev_feature = non_acad_dev_max = None
        st.info("Enter all staffing and enrollment figures to compute staffing indicators.")

    tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
        "Academic Content",
        "Staffing",
        "Physical Facilities",
        "Library",
        "Funding",
        "Research and Collaboration",
        "Tracer and Employers’ Rating",
    ])

    with tab1:
        academic = {
            "philosophy_objectives_defined": ask_scored_question("Philosophy and Objectives of the Programme", PHILOSOPHY_OPTIONS, "a0"),
            "curriculum_aligned_with_BMAS": ask_scored_question("Relationship between CCMAS and Curriculum", RELATIONSHIP_CCMAS_OPTIONS, "a1"),
            "innovative_courses_present": ask_scored_question("Innovation (Additional Courses)", INNOVATION_OPTIONS, "a2"),
            "curriculum_coverage_complete": ask_scored_question("Coverage of the Curriculum", COVERAGE_OPTIONS, "a3"),
            "admission_requirements_compliant": ask_scored_question("Admission Requirements into the Programme", ADMISSION_OPTIONS, "a4"),
            "academic_regulations_defined": ask_scored_question("Academic Regulations", ACADEMIC_REG_OPTIONS, "a5"),
            "tests_and_examinations_standardized": ask_scored_question("Standard of Tests and Examinations", TEST_EXAM_OPTIONS, "a6"),
            "evaluation_methods_clear": ask_scored_question("Evaluation of Students’ Work", EVALUATION_WORK_OPTIONS, "a7"),
            "degree_projects_adequate": ask_scored_question("Degree Projects", DEGREE_PROJECT_OPTIONS, "a8"),
            "practical_work_adequate": ask_scored_question("Practical Work", PRACTICAL_WORK_OPTIONS, "a9"),
            "student_course_evaluation_present": ask_scored_question("Students’ Course Evaluation", STUDENTS_COURSE_EVAL_OPTIONS, "a10"),
            "skills_acquisition_programme": ask_scored_question("Evaluation of Skills Acquisition", SKILLS_ACQUISITION_OPTIONS, "a11"),
            "external_examiner_system": ask_scored_question("External Examination System", EXTERNAL_EXAM_OPTIONS, "a12"),
            "internal_quality_assurance": ask_scored_question("Internal Quality Assurance System", IQA_OPTIONS, "a13"),
        }

    with tab2:
        staffing_noncomputed = {
            "staff_competence_verified": ask_scored_question("Competence of Teaching Staff", COMPETENCE_OPTIONS, "s1"),
            "administrative_support_available": ask_scored_question("Administration of College/School/Faculty/Department", ADMIN_OPTIONS, "s2"),
        }
        st.info("Calculated staffing items are automatically derived from the figures entered above.")

    with tab3:
        facilities = {
            "laboratory_space_adequate": ask_scored_question("Laboratories/Clinics/Studios/Farms/Museums (The existing space) is", LAB_SPACE_OPTIONS, "f1"),
            "laboratory_equipment_adequate": ask_scored_question("Laboratory Equipment (Meet the CCMAS Specifications by quality, quantity and functionality) up to", LAB_EQUIPMENT_OPTIONS, "f2"),
            "classroom_space_adequate": ask_scored_question("Classrooms/Lecture Theatres (The space available meets the space standards specified in the CCMAS) by", CLASSROOM_SPACE_OPTIONS, "f3"),
            "classroom_equipment_adequate": ask_scored_question("Classroom Equipment", CLASSROOM_EQUIPMENT_OPTIONS, "f4"),
            "office_accommodation_adequate": ask_scored_question("Office Accommodation", OFFICE_OPTIONS, "f5"),
            "safety_environment_present": ask_scored_question("Safety and Environment", SAFETY_OPTIONS, "f6"),
        }

    with tab4:
        library = {
            "library_holdings_adequate": ask_scored_question("Holdings", LIBRARY_HOLDINGS_OPTIONS, "l1"),
            "library_material_current": ask_scored_question("Currency of Holdings", CURRENCY_OPTIONS, "l2"),
            "e_library_subscription_available": ask_scored_question("Subscription to e-Books and e-Journals", SUBSCRIPTION_OPTIONS, "l3"),
            "e_library_access_good": ask_scored_question("Access to available e-Books and e-Journals", ACCESS_EBOOKS_OPTIONS, "l4"),
        }

    with tab5:
        funding_amount = st.number_input(
            "How much funding is allocated to the programme yearly? (₦)",
            min_value=0.0,
            value=None,
            placeholder="Enter amount"
        )
        funding_band, funding_score_value = classify_funding_amount(funding_amount)
        funding = {"programme_funding_adequate": funding_score_value}
        if funding_band is not None:
            st.info(f"Funding classification: {funding_band}")

    with tab6:
        research_collaboration = {
            "research_collaboration_active": ask_scored_question("Research and Collaboration", RESEARCH_COLLAB_OPTIONS, "r1"),
        }

    with tab7:
        has_graduated_students = st.selectbox(
            "Has the programme graduated students?",
            ["Yes", "No"],
            index=None,
            placeholder="Select an option",
            key="graduated_students"
        )

        tracer_options = None
        if has_graduated_students == "Yes":
            tracer_options = TRACER_OPTIONS
        elif has_graduated_students == "No":
            tracer_options = TRACER_NO_GRAD_OPTIONS

        tracer_rating = {
            "employer_rating_positive": ask_scored_question("Tracer and Employers’ Rating", tracer_options, "t1") if tracer_options else None,
        }

    academic_points = compute_section_points(academic, SECTION_MAX["Academic Content"])

    staffing_points = None
    if staffing_inputs_complete and all(v is not None for v in [
        ratio_score_raw, core_score_raw, mix_score_raw, phd_score_raw,
        non_teach_score_raw, acad_dev_score_raw, non_acad_dev_score_raw,
        staffing_noncomputed["staff_competence_verified"],
        staffing_noncomputed["administrative_support_available"],
    ]):
        total_staffing_raw = (
            ratio_score_raw + core_score_raw + mix_score_raw + phd_score_raw +
            non_teach_score_raw + acad_dev_score_raw + non_acad_dev_score_raw
        )
        total_staffing_max = (
            ratio_max + core_max + mix_max + phd_max +
            non_teach_max + acad_dev_max + non_acad_dev_max
        )
        computed_norm = total_staffing_raw / total_staffing_max if total_staffing_max else 0.0
        combined_staffing = {
            "computed_staffing_block": computed_norm,
            "staff_competence_verified": staffing_noncomputed["staff_competence_verified"],
            "administrative_support_available": staffing_noncomputed["administrative_support_available"],
        }
        staffing_points = compute_section_points(combined_staffing, SECTION_MAX["Staffing"])

    physical_facilities_points = compute_section_points(facilities, SECTION_MAX["Physical Facilities"])
    library_points = compute_section_points(library, SECTION_MAX["Library"])
    funding_points = compute_section_points(funding, SECTION_MAX["Funding"])
    research_collaboration_points = compute_section_points(research_collaboration, SECTION_MAX["Research and Collaboration"])
    tracer_points = compute_section_points(tracer_rating, SECTION_MAX["Tracer and Employers’ Rating"])

    if all(v is not None for v in [
        academic_points, staffing_points, physical_facilities_points,
        library_points, funding_points, research_collaboration_points,
        tracer_points
    ]):
        total_points = (
            academic_points +
            staffing_points +
            physical_facilities_points +
            library_points +
            funding_points +
            research_collaboration_points +
            tracer_points
        )
        programme_score = (total_points / TOTAL_MAX_SCORE) * 100
    else:
        total_points = None
        programme_score = None

    st.subheader("Predict Accreditation Status")

    required_basic = inputs_complete([institution, discipline, programme])
    required_all_sections = all(v is not None for v in [
        academic_points, staffing_points, physical_facilities_points,
        library_points, funding_points, research_collaboration_points,
        tracer_points, programme_score
    ])

    if st.button("Predict Accreditation Status", type="primary", width="stretch"):
        if not required_basic:
            st.error("Please complete the institutional information section.")
            st.stop()

        if not staffing_inputs_complete:
            st.error("Please complete all staffing and enrollment inputs.")
            st.stop()

        if not required_all_sections:
            st.error("Please complete all assessment sections before generating the result.")
            st.stop()

        input_data = {}
        input_data.update(academic)
        input_data.update(facilities)
        input_data.update(library)
        input_data.update(funding)
        input_data.update(research_collaboration)
        input_data.update(tracer_rating)
        input_data["staff_student_ratio_compliant"] = staff_ratio_feature
        input_data["proportion_core_staff_sufficient"] = core_staff_feature
        input_data["staff_rank_mix_balanced"] = staff_mix_feature
        input_data["academic_staff_qualification_high"] = phd_feature
        input_data["staff_competence_verified"] = staffing_noncomputed["staff_competence_verified"]
        input_data["administrative_support_available"] = staffing_noncomputed["administrative_support_available"]
        input_data["non_teaching_staff_adequate"] = non_teach_feature
        input_data["academic_staff_development_programme"] = acad_dev_feature
        input_data["non_academic_staff_development_programme"] = non_acad_dev_feature
        input_data["discipline"] = discipline

        input_df = pd.DataFrame([input_data])
        input_encoded = prepare_model_input(input_data, training_columns)
        model_status = predict_accreditation(model, label_encoder, input_encoded)
        probabilities = predict_probabilities(model, label_encoder, input_encoded)
        ml_confidence = probabilities.get(model_status, 0.0)
        predicted_status = determine_nuc_status(programme_score)

        st.subheader("Section Scores")

        row1 = st.columns(4)
        row2 = st.columns(4)

        with row1[0]:
            display_score_card("Academic Content", academic_points, SECTION_MAX["Academic Content"])
        with row1[1]:
            display_score_card("Staffing", staffing_points, SECTION_MAX["Staffing"])
        with row1[2]:
            display_score_card("Physical Facilities", physical_facilities_points, SECTION_MAX["Physical Facilities"])
        with row1[3]:
            display_score_card("Library", library_points, SECTION_MAX["Library"])

        with row2[0]:
            display_score_card("Funding", funding_points, SECTION_MAX["Funding"])
        with row2[1]:
            display_score_card("Research and Collaboration", research_collaboration_points, SECTION_MAX["Research and Collaboration"])
        with row2[2]:
            display_score_card("Tracer and Employers’ Rating", tracer_points, SECTION_MAX["Tracer and Employers’ Rating"])
        with row2[3]:
            st.markdown(
                f'<div class="score-card"><h4>Programme Score</h4><h2>{programme_score:.2f}%</h2><p>{total_points:.2f} / {TOTAL_MAX_SCORE}</p></div>',
                unsafe_allow_html=True
            )

        st.subheader("Score-Based Readiness Classification")
        st.success(predicted_status)
        st.subheader("Random Forest Prediction")
        st.info(f"{model_status} ({ml_confidence:.1%} confidence)")
        prob_cols = st.columns(3)
        for idx, status in enumerate(["FULL", "INTERIM", "DENIED"]):
            prob_cols[idx].metric(f"{status} Probability", f"{probabilities.get(status, 0):.1%}")
        st.caption(f"Model: {model_metadata['model_name']} | Version: {model_metadata['model_version']} | Training date: {model_metadata['training_date']}")
        if model_status != predicted_status:
            st.warning("The machine-learning prediction and score-based readiness classification differ. Review the highlighted criteria, supporting evidence and unresolved information before relying on the result.")
        else:
            st.success("The Random Forest prediction agrees with the score-based classification.")

        importance_df = get_feature_importance(model, training_columns)
        weak = []
        for col in input_df.columns:
            if col in importance_df["feature"].values:
                value = input_df[col].iloc[0]
                if value in [0.0, 0.5, 0.33, 0.25]:
                    impact = importance_df.loc[
                        importance_df["feature"] == col,
                        "importance"
                    ].values[0]
                    weak.append((col, value, impact))

        weak = sorted(weak, key=lambda x: x[2], reverse=True)

        if weak:
            weak_df = pd.DataFrame([
                {
                    "Feature": w[0].replace("_", " ").title(),
                    "Severity": "Critical" if w[1] in [0.0, 0.25] else "Moderate",
                    "Importance": round(float(w[2]), 4),
                }
                for w in weak[:10]
            ])
            st.subheader("Top High-Impact Weak Areas")
            st.dataframe(weak_df, width="stretch")
            top_weaknesses = "\n".join(
                [f"- {w[0].replace('_', ' ').title()}" for w in weak[:10]]
            )
        else:
            top_weaknesses = "- No major weaknesses were detected from the submitted responses."
            st.info("No major weaknesses were detected from the submitted responses.")

        with st.spinner("Generating accreditation report..."):
            report = generate_advisory_report(
                institution=institution,
                programme=programme,
                discipline=discipline,
                programme_score=programme_score,
                predicted_status=predicted_status,
                actual_ratio=actual_ratio,
                ratio_score_raw=ratio_score_raw,
                core_pct=core_pct,
                core_score_raw=core_score_raw,
                mix_score_raw=mix_score_raw,
                phd_pct=phd_pct,
                phd_score_raw=phd_score_raw,
                acad_dev_pct=acad_dev_pct,
                acad_dev_score_raw=acad_dev_score_raw,
                non_teach_score_raw=non_teach_score_raw,
                non_acad_dev_pct=non_acad_dev_pct,
                non_acad_dev_score_raw=non_acad_dev_score_raw,
                top_weaknesses=top_weaknesses,
            )

        st.subheader("Accreditation Report")
        st.markdown(report)

        save_assessment(
            user_id=st.session_state.user["id"],
            institution=institution,
            discipline=discipline,
            programme=programme,
            programme_score=float(programme_score),
            predicted_status=model_status,
            actual_ratio=float(actual_ratio),
            core_pct=float(core_pct),
            phd_pct=float(phd_pct),
            assessment_payload=json.dumps(input_data),
            report_text=report,
            assessment_type="standard_manual",
            rule_based_status=predicted_status,
            ml_confidence=float(ml_confidence),
            extraction_summary=None,
            evidence_coverage=100.0,
            total_points=float(total_points),
        )

        st.success("Assessment and report saved successfully.")
        st.download_button("Download Markdown Report", report, file_name="manual_readiness_report.md")
        st.download_button("Download JSON Assessment Record", json.dumps(input_data, indent=2), file_name="manual_assessment_record.json")
        st.download_button(
            "Download CSV Score Summary",
            pd.DataFrame([
                {"section": "Academic Content", "points": academic_points, "maximum": SECTION_MAX["Academic Content"]},
                {"section": "Staffing", "points": staffing_points, "maximum": SECTION_MAX["Staffing"]},
                {"section": "Physical Facilities", "points": physical_facilities_points, "maximum": SECTION_MAX["Physical Facilities"]},
                {"section": "Library", "points": library_points, "maximum": SECTION_MAX["Library"]},
                {"section": "Funding", "points": funding_points, "maximum": SECTION_MAX["Funding"]},
                {"section": "Research and Collaboration", "points": research_collaboration_points, "maximum": SECTION_MAX["Research and Collaboration"]},
                {"section": "Tracer and Employers’ Rating", "points": tracer_points, "maximum": SECTION_MAX["Tracer and Employers’ Rating"]},
            ]).to_csv(index=False),
            file_name="manual_score_summary.csv",
        )

# =========================================================
# SAVED ASSESSMENTS
# =========================================================
elif st.session_state.page == "Saved Assessments":
    page_intro(
        "Saved Assessments",
        "Review stored readiness records, model outputs, and evidence indicators for completed assessments.",
        None,
    )
    df_assessments = get_user_assessments(st.session_state.user["id"])

    if df_assessments.empty:
        render_empty_state(
            "empty_assessments.svg",
            "No saved assessments",
            "Completed manual and Computing document assessments will appear here after they are saved."
        )
    else:
        st.dataframe(df_assessments, width="stretch")

# =========================================================
# REPORT HISTORY
# =========================================================
elif st.session_state.page == "Report History":
    page_intro(
        "Report History",
        "Filter, inspect, and download generated accreditation-readiness reports.",
        None,
    )
    df_reports = get_user_report_history(st.session_state.user["id"])

    if df_reports.empty:
        render_empty_state(
            "empty_reports.svg",
            "No reports generated",
            "Advisory reports will appear here after a readiness prediction has been generated and saved."
        )
    else:
        f1, f2, f3 = st.columns(3)
        with f1:
            assessment_filter = st.selectbox("Assessment Type", ["All"] + sorted(df_reports["assessment_type"].dropna().unique().tolist()))
        with f2:
            programme_filter = st.selectbox("Programme", ["All"] + sorted(df_reports["programme_name"].dropna().unique().tolist()))
        with f3:
            status_filter = st.selectbox("Status", ["All"] + sorted(df_reports["rule_based_status"].dropna().unique().tolist()))

        filtered = df_reports.copy()
        if assessment_filter != "All":
            filtered = filtered[filtered["assessment_type"] == assessment_filter]
        if programme_filter != "All":
            filtered = filtered[filtered["programme_name"] == programme_filter]
        if status_filter != "All":
            filtered = filtered[filtered["rule_based_status"] == status_filter]

        for _, row in filtered.iterrows():
            title = f"{row['institution_name']} | {row['programme_name']} | {row['assessment_type']} | {row['rule_based_status']} | {row['created_at']}"
            with st.expander(title):
                st.write(f"**Programme Score:** {row['programme_score']:.2f}%")
                st.write(f"**Random Forest Prediction:** {row['predicted_status']} ({(row['ml_confidence'] or 0):.1%})")
                if row["evidence_coverage"] is not None:
                    st.write(f"**Evidence Coverage:** {row['evidence_coverage']:.2f}%")
                st.markdown(row["report_text"] or "No report text saved.")
                st.download_button(
                    "Download Report",
                    row["report_text"] or "",
                    file_name=f"assessment_{row['id']}_report.md",
                    key=f"download_report_{row['id']}",
                )

# =========================================================
# ADMIN DASHBOARD
# =========================================================
elif st.session_state.page == "Admin Dashboard" and st.session_state.user.get("role") == "admin":
    page_intro(
        "Admin Dashboard",
        "Monitor users, assessment activity, readiness distribution, and extraction review signals.",
        None,
    )

    users_df = get_all_users()
    assessments_df = get_all_assessments()

    d1, d2, d3, d4 = st.columns(4)
    with d1:
        st.metric("Total Users", len(users_df))
    with d2:
        st.metric("Total Assessments", len(assessments_df))
    with d3:
        if not assessments_df.empty:
            st.metric("Average Programme Score", f"{assessments_df['programme_score'].mean():.2f}%")
        else:
            st.metric("Average Programme Score", "0.00%")
    with d4:
        if not assessments_df.empty and "evidence_coverage" in assessments_df:
            st.metric("Average Evidence Coverage", f"{assessments_df['evidence_coverage'].fillna(100).mean():.2f}%")
        else:
            st.metric("Average Evidence Coverage", "0.00%")

    if not assessments_df.empty:
        a1, a2, a3 = st.columns(3)
        a1.metric("Standard Manual", int((assessments_df["assessment_type"] == "standard_manual").sum()))
        a2.metric("Computing Documents", int((assessments_df["assessment_type"] == "computing_document").sum()))
        disagreements = (assessments_df["rule_based_status"] != assessments_df["predicted_status"]).sum()
        a3.metric("Rule / Model Disagreements", int(disagreements))

        st.subheader("Status Distribution")
        status_distribution = assessments_df["rule_based_status"].value_counts()
        st.bar_chart(
            status_distribution.rename_axis("Status").reset_index(name="Count"),
            x="Status",
            y="Count",
            width="stretch",
        )

        conflicts = assessments_df[
            assessments_df["extraction_summary"].fillna("").str.contains("conflict", case=False, na=False)
        ]
        st.subheader("Extraction Conflicts")
        if conflicts.empty:
            st.caption("No extraction conflicts recorded.")
        else:
            st.dataframe(conflicts[["id", "institution_name", "programme_name", "extraction_summary", "created_at"]], width="stretch")

    st.subheader("Users")
    st.dataframe(users_df, width="stretch")

    st.subheader("All Assessments")
    st.dataframe(assessments_df, width="stretch")
