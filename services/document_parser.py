from io import BytesIO

import pandas as pd

from config.computing_requirements import DOCUMENT_CATEGORIES

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None


def _table_to_rows(table):
    return [[cell.text for cell in row.cells] for row in table.rows]


def _append_docx_section_text(text_blocks, section, paragraphs):
    text = "\n".join(paragraph.text for paragraph in paragraphs if paragraph.text.strip())
    if text.strip():
        text_blocks.append({"section": section, "text": text})


def read_uploaded_file(uploaded_file):
    suffix = uploaded_file.name.split(".")[-1].lower()
    data = uploaded_file.getvalue()
    return parse_document_bytes(data, uploaded_file.name, suffix)


def parse_document_bytes(data, file_name, suffix):
    warnings = []
    text_blocks = []
    tables = []

    if suffix == "pdf":
        if PdfReader is None:
            warnings.append("Install pypdf to extract PDF content.")
        else:
            reader = PdfReader(BytesIO(data))
            for index, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    label = "Title page" if index == 1 else f"Page {index}"
                    text_blocks.append({"section": label, "text": page_text})
            if not text_blocks:
                warnings.append("No meaningful extractable text found. OCR was not used.")
    elif suffix == "docx":
        if Document is None:
            warnings.append("Install python-docx to extract DOCX content.")
        else:
            document = Document(BytesIO(data))
            _append_docx_section_text(text_blocks, "Document body", document.paragraphs)
            for section_index, section in enumerate(document.sections, start=1):
                _append_docx_section_text(text_blocks, f"Header {section_index}", section.header.paragraphs)
                _append_docx_section_text(text_blocks, f"Footer {section_index}", section.footer.paragraphs)
                for table in section.header.tables:
                    tables.append(_table_to_rows(table))
                for table in section.footer.tables:
                    tables.append(_table_to_rows(table))
            for table in document.tables:
                tables.append(_table_to_rows(table))
    elif suffix == "xlsx":
        try:
            sheets = pd.read_excel(BytesIO(data), sheet_name=None)
            for sheet_name, sheet_df in sheets.items():
                tables.append(sheet_df.fillna("").astype(str).values.tolist())
                text_blocks.append({"section": f"Sheet: {sheet_name}", "text": sheet_df.to_csv(index=False)})
        except Exception as exc:
            warnings.append(f"Could not extract XLSX content: {exc}")
    elif suffix == "txt":
        text_blocks.append({"section": "Text file", "text": data.decode("utf-8", errors="ignore")})
    else:
        warnings.append("Unsupported file type.")

    combined = "\n".join(block["text"] for block in text_blocks).lower()
    detected = "unknown supporting evidence"
    for category, keywords in DOCUMENT_CATEGORIES.items():
        if any(keyword in combined for keyword in keywords):
            detected = category.replace("_", " ")
            break

    return {
        "file_name": file_name,
        "file_type": suffix,
        "text_blocks": text_blocks,
        "tables": tables,
        "detected_document_type": detected,
        "extraction_warnings": warnings,
        "ocr_used": False,
    }
